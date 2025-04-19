from langchain_community.document_loaders import Docx2txtLoader
from docx import Document  # 正确导入 python-docx
import io
import os
from typing import Optional



class DocProcessor:
    @staticmethod
    def load_doc(file_path):
        """优化的文档加载方法，直接使用python-docx提高效率"""
        try:
            # 直接使用python-docx读取文档
            doc = Document(file_path)
            
            # 使用列表推导式快速提取段落内容
            content = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
            
            # 改进错误处理
            if not content:
                print("警告：文档内容为空，返回默认值")
                return "请提供有效的文档内容"
            
            return content
            
        except Exception as e:
            print(f"加载文档时出现错误: {e}")
            return "加载文档时出现错误，请检查文档格式"
    
    @staticmethod
    def load_doc_stream(file_stream, filename: Optional[str] = None):
        """从流中直接读取文档，避免保存临时文件的开销"""
        try:
            # 直接从内存流读取
            doc = Document(io.BytesIO(file_stream))
            
            # 使用列表推导式快速提取段落内容
            content = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
            
            # 改进错误处理
            if not content:
                print("警告：文档内容为空，返回默认值")
                return "请提供有效的文档内容"
            
            return content
            
        except Exception as e:
            print(f"从流加载文档时出现错误: {e}")
            return "加载文档时出现错误，请检查文档格式"
    
    @staticmethod
    def save_doc(content: str, output_path: str):
        """保存结果到新文档"""
        doc = Document()  # 直接使用 Document
        
        # 按段落分割并添加到文档中
        for paragraph in content.split('\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph)
                
        doc.save(output_path)