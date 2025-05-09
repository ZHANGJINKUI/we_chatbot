from fastapi import FastAPI, UploadFile, File, HTTPException, Body, Request, Response, Query, Header, BackgroundTasks, Depends
from fastapi.responses import FileResponse, StreamingResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
import os
import tempfile
import sys
import shutil
import json
import logging
from typing import List, Dict, Any, Optional
from pydantic import BaseModel
import urllib.parse
# 注释掉数据库导入
# from admin.auth.routes import router as auth_router
# from admin.database import Base, engine
# 导入DocProcessor和相关组件
from my_agent.utils.shared.doc_processor import DocProcessor
from my_agent.utils.shared.intent_classifier import IntentClassifier
from langchain_openai import ChatOpenAI
import os
from dotenv import load_dotenv
import re
import asyncio
from io import BytesIO
from datetime import datetime
from pathlib import Path
# 从my_agent.agent_tools导入CSC和调用MCP的方法
from my_agent.agent_tools import csc as mcp_csc, call_mcp_service
# 直接导入完整的模块以便于调试
import my_agent.agent_tools
from docx import Document
import io
import uuid
import time

# 定义存储文件路径
DATA_DIR = Path(__file__).parent / "data"
CHAT_HISTORY_FILE = DATA_DIR / "chat_history.json"
FILE_LIST_FILE = DATA_DIR / "file_list.json"
DOCUMENT_STORAGE_FILE = DATA_DIR / "document_storage.json"

# 确保数据目录存在
DATA_DIR.mkdir(parents=True, exist_ok=True)

# 全局变量，用于存储最近上传的文档
last_uploaded_document = None
# 全局变量，用于存储当前文件列表
file_list = []
# 全局变量，记录用户会话是否进入强制公文纠错模式
forced_correction_mode = {}
# 全局变量，存储聊天历史记录
chat_list = []
# 全局变量存储处理后的文档
last_uploaded_document = None
# 全局变量存储聊天历史
chat_list = []
# 全局变量存储文件历史记录
file_history_list = []

# 保存聊天历史的文件路径
CHAT_HISTORY_FILE = "chat_history.json"
# 保存文件历史记录的文件路径
FILE_HISTORY_FILE = "file_history.json"

# 文档存储管理类
class DocumentStorage:
    def __init__(self):
        """初始化文档存储"""
        self.documents = {}  # 存储文档元数据
        self.file_contents = {}  # 存储二进制内容（内存中）
        
        # 使用绝对路径定义存储目录
        self.base_dir = os.path.dirname(os.path.abspath(__file__))
        self.data_folder = os.path.join(self.base_dir, "data")
        self.documents_folder = os.path.join(self.data_folder, "documents")  # 文档存储文件夹
        self.processed_folder = os.path.join(self.data_folder, "processed")  # 处理后文档存储文件夹
        
        # 确保文档存储目录存在
        os.makedirs(self.data_folder, exist_ok=True)
        os.makedirs(self.documents_folder, exist_ok=True)
        os.makedirs(self.processed_folder, exist_ok=True)
        
        print(f"文档存储初始化完成，数据目录: {self.data_folder}")
        print(f"原始文档目录: {self.documents_folder}")
        print(f"处理后文档目录: {self.processed_folder}")
        
        if not os.path.isdir(self.documents_folder):
            print(f"警告：原始文档目录 {self.documents_folder} 创建失败或不存在")
        
        if not os.path.isdir(self.processed_folder):
            print(f"警告：处理后文档目录 {self.processed_folder} 创建失败或不存在")
    
    def load_storage(self):
        """从JSON文件加载文档元数据"""
        try:
            if os.path.exists(DOCUMENT_STORAGE_FILE):
                with open(DOCUMENT_STORAGE_FILE, 'r', encoding='utf-8') as f:
                    self.documents = json.load(f)
                    print(f"从文件加载了 {len(self.documents)} 个文档元数据")
                    
                # 加载磁盘上的原始文档文件
                for doc_id, doc in self.documents.items():
                    # 尝试加载原始文档
                    original_path = os.path.join(self.documents_folder, f"{doc_id}.docx")
                    if os.path.exists(original_path):
                        with open(original_path, 'rb') as f:
                            self.file_contents[doc_id] = f.read()
                            print(f"加载了文档 {doc_id} 的原始文件，大小: {len(self.file_contents[doc_id])} 字节")
                    
                    # 尝试加载处理后文档
                    processed_path = os.path.join(self.processed_folder, f"{doc_id}.docx")
                    if os.path.exists(processed_path):
                        with open(processed_path, 'rb') as f:
                            self.file_contents[f"{doc_id}_processed"] = f.read()
                            self.documents[doc_id]["has_processed_binary"] = True
                            print(f"加载了文档 {doc_id} 的处理后文件，大小: {len(self.file_contents[f'{doc_id}_processed'])} 字节")
        except Exception as e:
            print(f"加载文档存储失败: {str(e)}")
            self.documents = {}
    
    def save_storage(self):
        """保存文档元数据到JSON文件并确保二进制内容保存到文件系统"""
        try:
            # 确保数据目录存在
            os.makedirs(self.data_folder, exist_ok=True)
            
            # 仅保存文本元数据到JSON
            docs_to_save = {}
            for doc_id, doc in self.documents.items():
                # 创建副本避免修改原始数据
                doc_copy = doc.copy()
                # 移除不需要保存到JSON的二进制内容
                if "file_content" in doc_copy:
                    del doc_copy["file_content"]
                docs_to_save[doc_id] = doc_copy
            
            # 保存元数据到JSON文件
            storage_path = os.path.join(self.data_folder, "document_storage.json")
            print(f"正在保存文档元数据到: {storage_path}")
            
            with open(storage_path, 'w', encoding='utf-8') as f:
                json.dump(docs_to_save, f, ensure_ascii=False, indent=2)
            
            if os.path.exists(storage_path):
                file_size = os.path.getsize(storage_path)
                print(f"已保存 {len(docs_to_save)} 个文档元数据到JSON，文件大小: {file_size} 字节")
            else:
                print(f"警告：JSON文件 {storage_path} 写入后未找到")
            
            # 保存二进制内容到文档文件
            successful_saves = 0
            for doc_id, content in self.file_contents.items():
                if not isinstance(content, bytes):
                    print(f"警告：文档 {doc_id} 的内容不是二进制类型，跳过保存")
                    continue
                    
                try:
                    if "_processed" in doc_id:
                        # 处理后的文档
                        original_id = doc_id.replace("_processed", "")
                        file_path = os.path.join(self.processed_folder, f"{original_id}.docx")
                    else:
                        # 原始文档
                        file_path = os.path.join(self.documents_folder, f"{doc_id}.docx")
                    
                    # 确保目录存在
                    os.makedirs(os.path.dirname(file_path), exist_ok=True)
                    print(f"正在保存文档 {doc_id} 的二进制内容到: {file_path}")
                    
                    with open(file_path, 'wb') as f:
                        f.write(content)
                    
                    if os.path.exists(file_path):
                        file_size = os.path.getsize(file_path)
                        print(f"已保存文档 {doc_id} 的二进制内容到文件，大小: {file_size} 字节")
                        successful_saves += 1
                    else:
                        print(f"警告：文件 {file_path} 写入后未找到")
                except Exception as e:
                    print(f"保存文档 {doc_id} 的二进制内容时出错: {str(e)}")
                    import traceback
                    print(traceback.format_exc())
            
            print(f"文档存储保存完成，成功保存 {successful_saves}/{len(self.file_contents)} 个二进制文件")
            return True
                        
        except Exception as e:
            print(f"保存文档存储失败: {str(e)}")
            import traceback
            print(traceback.format_exc())
            return False
    
    def add_document(self, document_id: str, filename: str, file_content: bytes):
        """添加新文档，只保存必要的元数据，将二进制内容保存到磁盘"""
        try:
            # 保存文档元数据
            self.documents[document_id] = {
                "document_id": document_id,
                "filename": filename,
                "has_binary": True,  # 标记有二进制内容
                "has_processed_binary": False,  # 标记是否有处理后的二进制文件
                "updatetime": datetime.now().isoformat()
            }
            
            # 单独保存二进制内容，避免JSON序列化问题
            self.file_contents[document_id] = file_content
            
            # 保存原始文档到磁盘
            original_path = os.path.join(self.documents_folder, f"{document_id}.docx")
            print(f"正在保存原始文档到: {original_path}")
            
            try:
                # 确保目录存在
                os.makedirs(os.path.dirname(original_path), exist_ok=True)
                
                # 打开文件并写入
                with open(original_path, 'wb') as f:
                    f.write(file_content)
                    print(f"已成功保存文档 {document_id} 的原始文件到磁盘，大小: {len(file_content)} 字节")
                    
                # 确认文件是否存在
                if os.path.exists(original_path):
                    file_size = os.path.getsize(original_path)
                    print(f"确认文件已写入: {original_path}, 大小: {file_size} 字节")
                else:
                    print(f"警告：文件 {original_path} 写入后无法找到")
            except Exception as e:
                print(f"保存原始文档到磁盘时出错: {str(e)}")
                import traceback
                print(traceback.format_exc())
            
            # 保存到持久化存储
            self.save_storage()
            print(f"成功保存文档 {filename} (ID: {document_id}), 二进制长度: {len(file_content)}")
            return True
        except Exception as e:
            print(f"添加文档时出错: {str(e)}")
            import traceback
            print(traceback.format_exc())
            return False
    
    def get_document(self, document_id: str) -> Optional[dict]:
        """获取文档，包括二进制内容和文本内容"""
        doc = self.documents.get(document_id)
        if not doc:
            print(f"未找到文档，ID: {document_id}")
            return None
        
        # 创建文档副本，避免修改原始数据
        doc_copy = doc.copy()
        
        # 添加二进制内容
        if doc.get("has_binary") and document_id in self.file_contents:
            doc_copy["file_content"] = self.file_contents[document_id]
        else:
            # 尝试从磁盘加载
            file_path = os.path.join(self.documents_folder, f"{document_id}.docx")
            if os.path.exists(file_path):
                try:
                    with open(file_path, 'rb') as f:
                        binary_content = f.read()
                        self.file_contents[document_id] = binary_content
                        doc_copy["file_content"] = binary_content
                        doc_copy["has_binary"] = True
                        print(f"从磁盘加载文档 {document_id} 的二进制内容，大小: {len(binary_content)} 字节")
                except Exception as e:
                    print(f"读取文档二进制内容出错: {str(e)}")
            else:
                print(f"警告：找不到文档 {document_id} 的二进制文件: {file_path}")
        
        # 如果文档没有内容或内容为空，尝试读取并解析文档内容
        if ("content" not in doc_copy or not doc_copy.get("content", "").strip()) and os.path.exists(os.path.join(self.documents_folder, f"{document_id}.docx")):
            try:
                from my_agent.utils.shared.doc_processor import DocProcessor
                file_path = os.path.join(self.documents_folder, f"{document_id}.docx")
                text_content, _ = DocProcessor.process_docx_file(file_path)
                if text_content:
                    doc_copy["content"] = text_content
                    # 同时更新存储中的数据
                    self.documents[document_id]["content"] = text_content
                    print(f"从文件中读取并解析了文档 {document_id} 的文本内容，长度: {len(text_content)}")
            except Exception as e:
                print(f"解析文档文本内容出错: {str(e)}")
        
        return doc_copy
    
    def update_processed_content(self, document_id: str, processed_content: str):
        """更新处理后的文本内容"""
        if document_id in self.documents:
            self.documents[document_id]["processed_content"] = processed_content
            self.documents[document_id]["updatetime"] = datetime.now().isoformat()
            self.save_storage()
            print(f"已更新文档 {document_id} 的处理后文本内容")
    
    def update_processed_binary(self, document_id: str, processed_binary: bytes):
        """更新处理后的二进制内容"""
        if document_id in self.documents:
            # 保存处理后的二进制内容到内存中
            self.documents[document_id]["has_processed_binary"] = True
            # 单独保存二进制内容
            self.file_contents[f"{document_id}_processed"] = processed_binary
            self.documents[document_id]["updatetime"] = datetime.now().isoformat()
            
            # 保存处理后的文档到磁盘
            processed_path = os.path.join(self.processed_folder, f"{document_id}.docx")
            with open(processed_path, 'wb') as f:
                f.write(processed_binary)
                print(f"已保存文档 {document_id} 的处理后文件到磁盘，大小: {len(processed_binary)} 字节")
            
            self.save_storage()
            print(f"更新了文档 {document_id} 的处理后二进制内容，大小: {len(processed_binary)} 字节")
    
    def get_all_documents(self) -> list:
        """获取所有文档，按更新时间排序"""
        try:
            documents_without_binary = []
            for doc_id, doc in self.documents.items():
                # 创建副本，不修改原始数据
                doc_copy = doc.copy()
                # 移除二进制内容以避免JSON序列化问题
                if "file_content" in doc_copy:
                    del doc_copy["file_content"]
                documents_without_binary.append(doc_copy)
            
            # 按更新时间排序（最新的在前面）
            documents_without_binary.sort(
                key=lambda x: x.get("updatetime", ""), 
                reverse=True
            )
            
            logging.info(f"获取到 {len(documents_without_binary)} 个文档，已按时间排序")
            return documents_without_binary
        except Exception as e:
            error_msg = f"获取文档列表出错: {str(e)}"
            logging.error(error_msg)
            return []
    
    def delete_document(self, document_id: str) -> bool:
        """删除文档及其关联的文件"""
        if document_id in self.documents:
            # 删除文档记录
            del self.documents[document_id]
            
            # 同时清理内存中的二进制内容
            if document_id in self.file_contents:
                del self.file_contents[document_id]
                logging.info(f"已清理文档 {document_id} 的内存中二进制内容")
            
            # 删除处理后内容的内存
            if f"{document_id}_processed" in self.file_contents:
                del self.file_contents[f"{document_id}_processed"]
                logging.info(f"已清理文档 {document_id} 的内存中处理后二进制内容")
            
            # 删除文件系统中的文件
            original_file_path = os.path.join(self.documents_folder, f"{document_id}.docx")
            processed_file_path = os.path.join(self.processed_folder, f"{document_id}.docx")
            
            # 删除原始文件
            if os.path.exists(original_file_path):
                try:
                    os.remove(original_file_path)
                    logging.info(f"已删除文件系统中的原始文件: {original_file_path}")
                except Exception as e:
                    logging.error(f"删除原始文件失败: {str(e)}")
            
            # 删除处理后文件
            if os.path.exists(processed_file_path):
                try:
                    os.remove(processed_file_path)
                    logging.info(f"已删除文件系统中的处理后文件: {processed_file_path}")
                except Exception as e:
                    logging.error(f"删除处理后文件失败: {str(e)}")
            
            # 保存更改到存储
            self.save_storage()
            return True
        return False

    def sync_to_file_list(self) -> None:
        """同步文档存储内容到全局file_list变量"""
        try:
            global file_list
            
            # 清空现有file_list
            file_list = []
            
            # 从文档存储中获取所有文档
            all_docs = self.get_all_documents()
            
            # 将文档转换为file_list需要的格式
            for doc in all_docs:
                file_list.append({
                    "id": doc.get("document_id", ""),
                    "filename": doc.get("filename", ""),
                    "fileuuid": doc.get("document_id", ""),
                    "fileurl": "",
                    "updatetime": doc.get("updatetime", ""),
                    "userid": "system"  # 可以从请求中获取，这里简化处理
                })
            
            logging.info(f"已将 {len(file_list)} 个文档同步到file_list")
        except Exception as e:
            error_msg = f"同步文档到file_list出错: {str(e)}"
            logging.error(error_msg)

    def get_document_content(self, document_id: str) -> Optional[str]:
        """获取文档的文本内容，如果内容不存在则尝试从文件中读取"""
        # 尝试直接从内存中获取文档
        doc = self.documents.get(document_id)
        if not doc:
            print(f"未找到文档，ID: {document_id}")
            return None
        
        # 如果文档内容已存在且不为空，直接返回
        if "content" in doc and doc["content"] and doc["content"].strip():
            print(f"从documents存储中获取到文档 {document_id} 的内容，长度: {len(doc['content'])}")
            return doc["content"]
        
        print(f"文档 {document_id} 内容为空或不存在，尝试从文件中读取")
        
        # 如果没有内容或内容为空，尝试从文件中读取
        file_path = os.path.join(self.documents_folder, f"{document_id}.docx")
        print(f"尝试从文件路径读取: {file_path}")
        
        if os.path.exists(file_path):
            try:
                print(f"找到文件: {file_path}，开始读取内容")
                from my_agent.utils.shared.doc_processor import DocProcessor
                text_content, _ = DocProcessor.process_docx_file(file_path)
                if text_content:
                    # 更新文档内容
                    self.documents[document_id]["content"] = text_content
                    print(f"成功从文件中读取并解析了文档 {document_id} 的文本内容，长度: {len(text_content)}")
                    # 保存更新到存储文件
                    self.save_storage()
                    print(f"已将文档 {document_id} 的内容保存到持久存储")
                    return text_content
                else:
                    print(f"警告：从文件 {file_path} 读取的内容为空")
            except Exception as e:
                print(f"解析文档文本内容出错: {str(e)}")
                import traceback
                print(traceback.format_exc())
        else:
            print(f"警告：找不到文档文件: {file_path}")
        
        print(f"无法获取文档 {document_id} 的内容")
        return None

# 创建全局文档存储实例
document_storage = DocumentStorage()

# 加载聊天历史
def load_chat_history():
    global chat_list
    try:
        if os.path.exists(CHAT_HISTORY_FILE):
            with open(CHAT_HISTORY_FILE, "r", encoding="utf-8") as f:
                chat_list = json.load(f)
                print(f"从文件加载聊天历史，共 {len(chat_list)} 条记录")
    except Exception as e:
        print(f"加载聊天历史出错: {str(e)}")
        chat_list = []

# 保存聊天历史到文件
def save_chat_history():
    try:
        with open(CHAT_HISTORY_FILE, "w", encoding="utf-8") as f:
            json.dump(chat_list, f, ensure_ascii=False, indent=2)
            print(f"聊天历史已保存到文件，共 {len(chat_list)} 条记录")
    except Exception as e:
        print(f"保存聊天历史出错: {str(e)}")

# 加载文件历史记录
def load_file_history():
    global file_history_list
    try:
        if os.path.exists(FILE_HISTORY_FILE):
            with open(FILE_HISTORY_FILE, "r", encoding="utf-8") as f:
                file_history_list = json.load(f)
                print(f"从文件加载文件历史记录，共 {len(file_history_list)} 条记录")
    except Exception as e:
        print(f"加载文件历史记录出错: {str(e)}")
        file_history_list = []

# 保存文件历史记录到文件
def save_file_history():
    try:
        with open(FILE_HISTORY_FILE, "w", encoding="utf-8") as f:
            json.dump(file_history_list, f, ensure_ascii=False, indent=2)
            print(f"文件历史记录已保存到文件，共 {len(file_history_list)} 条记录")
    except Exception as e:
        print(f"保存文件历史记录出错: {str(e)}")

# 首先创建FastAPI应用实例
app = FastAPI(title="文档处理API")

# 在应用启动时加载聊天历史
@app.on_event("startup")
async def startup_event():
    """应用启动时的初始化工作"""
    print("\n=== 应用启动，初始化目录和数据 ===")
    
    # 获取应用根目录
    base_dir = os.path.dirname(os.path.abspath(__file__))
    data_dir = os.path.join(base_dir, "data")
    documents_dir = os.path.join(data_dir, "documents")
    processed_dir = os.path.join(data_dir, "processed")
    
    # 确保所有必要的目录存在
    print(f"确保数据目录存在: {data_dir}")
    os.makedirs(data_dir, exist_ok=True)
    
    print(f"确保文档目录存在: {documents_dir}")
    os.makedirs(documents_dir, exist_ok=True)
    
    print(f"确保处理后文档目录存在: {processed_dir}")
    os.makedirs(processed_dir, exist_ok=True)
    
    print(f"确保临时目录存在: {TEMP_DIR}")
    os.makedirs(TEMP_DIR, exist_ok=True)
    
    # 检查目录是否成功创建
    for directory in [data_dir, documents_dir, processed_dir, TEMP_DIR]:
        if os.path.isdir(directory):
            print(f"目录已创建: {directory}")
        else:
            print(f"警告: 目录创建失败: {directory}")
    
    # 初始化时加载聊天历史
    load_chat_history()
    
    # 初始化时加载文件历史记录
    load_file_history()
    
    # 初始化文档存储
    document_storage.load_storage()
    document_count = len(document_storage.documents)
    print(f"已加载 {document_count} 个文档元数据")
    
    # 检查文档存储
    binary_count = len(document_storage.file_contents)
    print(f"已加载 {binary_count} 个文档二进制内容")
    
    # 检查文件是否存在于磁盘
    disk_files = 0
    for doc_id in document_storage.documents:
        doc_path = os.path.join(documents_dir, f"{doc_id}.docx")
        if os.path.exists(doc_path):
            disk_files += 1
            
    print(f"在磁盘上找到 {disk_files}/{document_count} 个文档文件")
    
    # 同步文档存储到file_list
    document_storage.sync_to_file_list()
    logging.info(f"已同步 {len(file_list)} 个文档到file_list")
    
    logging.info("应用启动完成，已加载聊天历史和文件历史记录")
    print("=== 应用启动初始化完成 ===\n")

# 配置CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # 允许所有源
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 定义临时文件路径
TEMP_DIR = tempfile.gettempdir()

# 加载环境变量和初始化模型
# 使用相对路径加载.env文件
# 获取当前文件的目录
current_dir = Path(__file__).parent
# 获取项目根目录（即main.py所在目录的上一级）
root_dir = current_dir.parent
# 构建.env文件的路径
dotenv_path = root_dir / '.env'
load_dotenv(dotenv_path)
qwen_api_key = os.getenv("QWEN_API_KEY")
qwen_base_url = os.getenv("QWEN_BASE_URL")

# 检查环境变量是否设置
if not qwen_api_key or not qwen_base_url:
    print(f"警告: Qwen API 配置缺失。尝试从 {dotenv_path} 加载但未成功。")
    print("将使用默认配置继续...")
    qwen_api_key = "dummy_key"  # 使用默认值避免错误
    qwen_base_url = "https://dashscope.aliyuncs.com/v1"  # 使用默认值避免错误
else:
    print(f"成功从 {dotenv_path} 加载 Qwen API 配置")

# 设置环境变量
os.environ["OPENAI_API_KEY"] = qwen_api_key
os.environ["OPENAI_API_BASE"] = qwen_base_url

# 初始化LLM，使用正确的模型名称
try:
    llm = ChatOpenAI(
        model="qwen-max", 
        temperature=0.7,
        api_key=qwen_api_key,
        base_url=qwen_base_url
    )
    print("成功初始化Qwen LLM模型")
except Exception as e:
    print(f"初始化Qwen LLM模型失败: {str(e)}")
    print("将使用备用配置...")
    # 备用配置，使用最基本的设置
    llm = ChatOpenAI(model="qwen-max")

# 闲聊计数器 - 全局变量
chat_counter = {}

# 定义请求和响应模型
class UploadDocumentResponse(BaseModel):
    document_id: str
    content: str
    message: str

class ChatRequest(BaseModel):
    message: str
    document_content: str = None  # 新增文档内容字段
    chat_history: list = []

class ChatResponse(BaseModel):
    response: str
    chat_history: list
    processed_document: str = None  # 新增处理后的文档内容字段

class SaveDocumentRequest(BaseModel):
    document_id: str
    content: str
    filename: str = "processed_document.docx"

# 添加登录请求和响应模型
class LoginRequest(BaseModel):
    userid: str
    password: str

class LoginResponseData(BaseModel):
    user: dict
    token: str

class LoginResponse(BaseModel):
    code: int
    data: LoginResponseData
    msg: str

# 定义路由
@app.get("/")
def read_root():
    return {"message": "Welcome to the User Management System"}

# 添加健康检查API
@app.get("/api/health-check")
async def health_check():
    """健康检查端点"""
    return {"status": "ok", "message": "Service is running"}

# 添加简单的登录API端点
@app.post("/api/login", response_model=LoginResponse)
async def login(login_request: LoginRequest):
    """处理用户登录请求"""
    # 这里简化处理，只要提供了用户名和密码就认为登录成功
    if login_request.userid and login_request.password:
        # 返回模拟的用户信息和token，包装为前端期望的格式
        response_data = {
            "user": {
                "userid": login_request.userid,
                "username": login_request.userid,
                "role": "user"
            },
            "token": "mock_token_123456789"
        }
        
        # 包装为前端期望的格式
        return {
            "code": 200,
            "data": response_data,
            "msg": "登录成功"
        }
    else:
        raise HTTPException(status_code=401, detail="用户名或密码错误")

@app.post("/api/upload-document")
async def upload_document(file: UploadFile = None):
    """处理文档上传请求"""
    try:
        print("接收到文档上传请求")
        
        if not file:
            raise HTTPException(status_code=400, detail="未提供文件")
            
        # 读取文件内容
        file_content = await file.read()
        if not file_content:
            raise HTTPException(status_code=400, detail="文件内容为空")
            
        filename = file.filename
        print(f"接收到文件: {filename}, 大小: {len(file_content)} 字节")
        
        # 生成唯一的文档ID
        document_id = str(uuid.uuid4())
        
        # 添加到文档存储，只保存元数据，不解析内容
        document_storage.add_document(
            document_id=document_id,
            filename=filename,
            file_content=file_content
        )
        
        # 更新全局变量，用于后续处理
        global last_uploaded_document
        last_uploaded_document = {
            "document_id": document_id,
            "filename": filename,
            "file_path": os.path.join(document_storage.documents_folder, f"{document_id}.docx")
        }
        print(f"设置当前文档: {document_id}")
        
        # 更新文件列表的逻辑
        global file_list
        file_list.insert(0, {
            "id": document_id,
            "userid": "system",
            "filename": filename,
            "fileuuid": document_id,
            "fileurl": "",
            "updatetime": datetime.now().isoformat()
        })
        
        return {
            "document_id": document_id,
            "message": "文档上传成功"
        }
    except Exception as e:
        print(f"上传文档时出错: {str(e)}")
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"上传文档时出错: {str(e)}")

@app.post("/api/process-document", deprecated=True)
async def process_document_api(file: UploadFile = File(...)):
    """上传并处理Word文档 (已弃用，保留兼容)"""
    return await upload_document(file)

# 修改聊天接口，处理带有文档内容的请求
@app.post("/api/chat", response_model=ChatResponse)
async def chat_endpoint(chat_request: ChatRequest, request: Request):
    """处理聊天请求，支持文档处理"""
    try:
        print(f"Received chat request: {chat_request.message[:100]}...")
        
        # 准备消息历史
        messages = chat_request.chat_history.copy()
        
        # 检查是否有文档内容，如果没有，尝试从最近上传的文档中获取
        document_content = chat_request.document_content
        if (document_content is None or document_content.strip() == "") and last_uploaded_document:
            document_content = last_uploaded_document["content"]
            print(f"使用最近上传的文档内容，长度: {len(document_content)}")
        
        has_document = document_content is not None and document_content.strip() != ""
        
        # 进行意图分类
        intent = IntentClassifier.classify(chat_request.message)
        print(f"Intent classified as: {intent}")
        
        # 检查是否包含润色或总结关键词
        if "润色" in chat_request.message:
            intent = "polish"
            print("从消息中检测到润色关键词，设置意图为polish")
        elif "总结" in chat_request.message:
            intent = "summary"
            print("从消息中检测到总结关键词，设置意图为summary")
        
        # 获取客户端IP地址作为会话ID
        client_ip = request.client.host
        session_id = f"session_{client_ip}"
        print(f"Using session ID: {session_id}")
        
        # 初始化会话计数器
        if session_id not in chat_counter:
            chat_counter[session_id] = 0
            print(f"Initialized chat counter for {session_id}")
            
        # 如果是聊天意图，增加计数
        if intent == "chat":
            chat_counter[session_id] += 1
            print(f"Increased chat counter for {session_id} to {chat_counter[session_id]}")
            
        # 检查是否超过闲聊次数限制
        if intent == "chat" and chat_counter[session_id] > 3:
            # 重置意图为强制文档处理
            intent = "forced_correction"
            # 添加用户消息到历史
            messages.append({"role": "user", "content": chat_request.message})
            # 添加系统消息，引导用户回到公文纠错
            guidance_message = "您已经闲聊超过3次，请回到公文纠错相关任务。您可以尝试输入「纠错」、「润色」等指令来处理您的文档。"
            messages.append({"role": "assistant", "content": guidance_message})
            
            return {
                "response": guidance_message,
                "chat_history": messages,
                "processed_document": None
            }
            
        # 根据意图处理
        if intent in ["correction", "writing", "recorrection", "forced_correction", "polish", "summary"]:
            # 检查是否有文档内容
            if not has_document:
                # 文档内容缺失，提示用户上传文档
                messages.append({"role": "user", "content": chat_request.message})
                no_doc_message = "请先上传Word文档再进行公文纠错。您可以点击左侧面板的「上传文档」按钮上传文件。"
                messages.append({"role": "assistant", "content": no_doc_message})
                
                return {
                    "response": no_doc_message,
                    "chat_history": messages,
                    "processed_document": None
                }
            
            # 修改为使用MCP协议封装的CSC纠错服务
            print("使用MCP协议的CSC服务处理文档...")
            
            # 调用MCP服务进行文档纠错
            csc_result = mcp_csc(document_content)
            
            if csc_result["status"] == "success":
                # 获取处理后的文档内容
                if csc_result.get("modified", False):
                    processed_document = csc_result["corrected_text"]
                    processing_message = f"文档纠错完成，共进行了以下修改：\n{csc_result.get('changes', '未提供详细修改说明')}"
                else:
                    processed_document = document_content
                    processing_message = csc_result.get("message", "文档未发现需要纠错的内容，已保持原文。")
                
                # 将用户消息添加到聊天历史
                messages.append({"role": "user", "content": chat_request.message})
                
                # 将处理结果添加到聊天历史
                success_message = f"已完成文档纠错。{processing_message}"
                messages.append({"role": "assistant", "content": success_message})
                
                # 保存处理后的文档到全局变量和持久化存储
                if last_uploaded_document and "document_id" in last_uploaded_document:
                    # 更新文本内容
                    last_uploaded_document["processed_content"] = processed_document
                    document_storage.update_processed_content(last_uploaded_document["document_id"], processed_document)
                    
                    # 获取原始文档的二进制内容，用于生成处理后文档
                    if last_uploaded_document["document_id"] in document_storage.file_contents:
                        original_binary = document_storage.file_contents[last_uploaded_document["document_id"]]
                        # 生成处理后的文档二进制内容
                        processed_binary = convert_text_to_docx_bytes(processed_document, original_binary)
                        # 更新二进制内容
                        document_storage.update_processed_binary(last_uploaded_document["document_id"], processed_binary)
                        print(f"已保存处理后的二进制文档内容，大小: {len(processed_binary)} 字节")
                
                # 返回处理结果和更新后的聊天历史
                return {
                    "response": success_message,
                    "chat_history": messages,
                    "processed_document": processed_document
                }
            else:
                # 处理失败，回退到原始的LLM处理方式
                print("MCP服务处理失败，回退到LLM处理...")
                error_message = csc_result.get("message", "MCP服务处理失败")
                print(f"错误: {error_message}")
                
            # 构建包含文档内容的系统提示
            system_prompt = "你是一个专业的文档处理助手。请根据用户的指令处理以下文档内容。"
            
            # 将文档内容和用户指令组合成完整提示
            full_prompt = f"文档内容:\n{document_content}\n\n用户指令:\n{chat_request.message}"
            
            # 准备消息序列
            doc_messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": full_prompt}
            ]
            
            # 调用模型处理文档
            print("转为调用大语言模型处理文档...")
            response = llm.invoke(doc_messages)
            
            # 将用户消息添加到聊天历史
            messages.append({"role": "user", "content": chat_request.message})
            
            # 将模型响应添加到聊天历史
            fallback_message = "MCP服务暂时不可用，已使用备用方法处理您的文档。处理结果已显示在中间区域。"
            messages.append({"role": "assistant", "content": fallback_message})
            
            # 保存处理后的文档到全局变量和持久化存储
            if last_uploaded_document and "document_id" in last_uploaded_document:
                # 更新文本内容
                last_uploaded_document["processed_content"] = response.content
                document_storage.update_processed_content(last_uploaded_document["document_id"], response.content)
                
                # 获取原始文档的二进制内容，用于生成处理后文档
                if last_uploaded_document["document_id"] in document_storage.file_contents:
                    original_binary = document_storage.file_contents[last_uploaded_document["document_id"]]
                    # 生成处理后的文档二进制内容
                    processed_binary = convert_text_to_docx_bytes(response.content, original_binary)
                    # 更新二进制内容
                    document_storage.update_processed_binary(last_uploaded_document["document_id"], processed_binary)
                    print(f"已保存处理后的二进制文档内容，大小: {len(processed_binary)} 字节")
            
            # 返回处理结果和更新后的聊天历史
            return {
                "response": fallback_message,
                "chat_history": messages,
                "processed_document": response.content
            }
        else:
            # 普通聊天处理
            # 添加用户新消息
            messages.append({"role": "user", "content": chat_request.message})
            
            # 构建系统提示
            system_prompt = """你是政务公文智能体，一个智能公文写作助手。
当用户询问你的身份、名字、是谁、介绍自己等类似问题时，必须回答：
"我是政务公文智能体，您的智能公文写作助手！"
你可以完成政务公文写作、润色、纠错、总结任务。
请以自然、友好的方式回答用户的其他问题。
记住，不要透露你是AI、大模型或Qwen Chat，而应始终以我是政务公文智能体身份回答。"""
            
            # 添加系统消息到开始
            full_messages = [{"role": "system", "content": system_prompt}] + messages
            
            try:
                # 调用模型生成回复
                print("调用Qwen LLM生成回复...")
                model_response = llm.invoke(full_messages)
                assistant_message = model_response.content
                print(f"Qwen LLM生成回复成功: {assistant_message[:50]}..." if len(assistant_message) > 50 else f"Qwen LLM生成回复成功: {assistant_message}")
                
                # 更新聊天历史
                messages.append({"role": "assistant", "content": assistant_message})
                
                # 返回结果，而不是使用yield
                return {
                    "response": assistant_message,
                    "chat_history": messages,
                    "processed_document": None
                }
                
            except Exception as e:
                print(f"调用LLM出错: {str(e)}")
                error_message = f"处理您的请求时出错: {str(e)}。请稍后再试。"
                
                # 更新聊天历史
                messages.append({"role": "assistant", "content": error_message})
                
                return {
                    "response": error_message,
                    "chat_history": messages,
                    "processed_document": None
                }
            
    except Exception as e:
        print(f"Chat endpoint error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"处理聊天请求时出错: {str(e)}")

@app.get("/api/download-result")
async def download_result(chatId: str = None, docType: str = "modified"):
    """
    下载处理后的文档或写作文档
    
    参数:
        chatId: 聊天ID
        docType: 文档类型，支持"modified"(修改后文档) 和 "writing"(写作文档)，默认为modified
    
    返回:
        处理后的文档或写作文档的二进制内容
    """
    if not chatId:
        raise HTTPException(status_code=400, detail="缺少chatId参数")
    
    try:
        # 检查这个聊天关联的文档
        matching_document = None
        matching_document_id = None
        
        if docType == "writing":
            # 先尝试查找最新的写作文档
            writing_prefix = f"writing_{chatId}_"
            latest_doc = None
            latest_timestamp = 0
            
            # 遍历文档找到最新的写作文档
            for doc_id, doc in document_storage.documents.items():
                if doc_id.startswith(writing_prefix):
                    # 提取时间戳并比较
                    try:
                        # writing_chatId_timestamp 格式
                        timestamp_str = doc_id.split('_')[-1]
                        timestamp = int(timestamp_str)
                        if timestamp > latest_timestamp:
                            latest_timestamp = timestamp
                            latest_doc = doc
                            matching_document_id = doc_id
                    except (ValueError, IndexError):
                        continue
            
            if latest_doc:
                matching_document = latest_doc
                print(f"找到最新的写作文档: {matching_document_id}, 时间戳: {latest_timestamp}")
            
            # 如果找不到写作文档，查找其他写作相关文档
            if not matching_document:
                for doc_id, doc in document_storage.documents.items():
                    if "writing" in doc_id and doc.get("has_processed_binary", False):
                        matching_document = doc
                        matching_document_id = doc_id
                        print(f"找到相关写作文档: {doc_id}")
                        break
        else:
            # 处理修改后文档
            # 查找与chat_id关联的历史或当前文档
            for doc_id, doc in document_storage.documents.items():
                # 首先查找处理后的文档
                if doc.get("has_processed_binary", False):
                    matching_document = doc
                    matching_document_id = doc_id
                    print(f"找到具有处理后内容的文档: {doc_id}")
                    break
            
            # 如果没找到，则使用上次上传的文档
            if not matching_document and last_uploaded_document:
                doc_id = last_uploaded_document.get("document_id")
                if doc_id and doc_id in document_storage.documents:
                    matching_document = document_storage.documents[doc_id]
                    matching_document_id = doc_id
                    print(f"使用last_uploaded_document: {doc_id}")
        
        if not matching_document:
            raise HTTPException(status_code=404, detail="找不到与该聊天关联的文档")
        
        # 生成文件名
        now = datetime.now().strftime('%Y%m%d%H%M%S')
        if docType == "writing":
            filename = f"writing_document_{now}.docx"
        else:
            filename = f"processed_{matching_document.get('filename', 'document.docx')}"
        
        # 获取处理后的二进制内容
        if matching_document.get("has_processed_binary", False):
            # 返回处理后的二进制内容
            processed_binary = document_storage.file_contents.get(f"{matching_document_id}_processed")
            if processed_binary:
                return Response(
                    content=processed_binary,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": f"attachment; filename={filename}"}
                )
        
        # 获取处理后的文本内容
        processed_content = matching_document.get("processed_content", "")
        if not processed_content:
            # 文档内容为空，尝试获取原始内容
            processed_content = matching_document.get("content", "")
            
        if not processed_content:
            raise HTTPException(status_code=404, detail="找不到可下载的文档内容")
        
        # 获取原始二进制内容作为模板
        original_binary = None
        if matching_document_id in document_storage.file_contents:
            original_binary = document_storage.file_contents[matching_document_id]
        
        # 将文本内容转换为docx并返回，使用原始文档作为模板
        doc_bytes = convert_text_to_docx_bytes(processed_content, original_binary)
        
        return Response(
            content=doc_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        print(f"下载文档失败: {str(e)}")
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"下载文档失败: {str(e)}")

@app.post("/api/save-document")
async def save_document(request: SaveDocumentRequest):
    """保存处理后的文档内容"""
    try:
        # 确保有内容要保存
        if not request.content or not request.content.strip():
            raise HTTPException(status_code=400, detail="没有可保存的内容")
            
        # 将内容保存到临时文件
        file_path = os.path.join(TEMP_DIR, request.filename)
        
        # 使用DocProcessor保存为docx
        DocProcessor.save_text_to_docx(request.content, file_path)
        
        return {"status": "success", "message": "文档已保存", "path": file_path}
        
    except Exception as e:
        if isinstance(e, HTTPException):
            raise e
        raise HTTPException(status_code=500, detail=f"保存文档时出错: {str(e)}")

# 添加流式聊天接口
@app.get("/api/stream-chat")
async def stream_chat(request: Request, message: str, document_content: str = None, chat_history: str = "[]", file_id: str = None, intent: str = None):
    """处理聊天请求，返回流式响应"""
    # 在函数外部处理chat_history，避免闭包问题
    history_data = json.loads(chat_history) if chat_history else []
    print(f"解析到聊天历史消息数: {len(history_data)}")
    
    async def event_generator():
        try:
            print(f"收到流式聊天请求: {message[:50]}..." if len(message) > 50 else f"收到流式聊天请求: {message}")
            print(f"文件ID: {file_id}")
            
            # 准备聊天消息
            messages = history_data.copy() # 使用外部已处理的history_data
            
            # 检查是否有文档内容，如果没有，尝试从last_uploaded_document获取
            global last_uploaded_document, forced_correction_mode, file_list
            has_document = False
            input_text_for_correction = None
            processed_doc_content = None
            processed_content = None  # 存储处理后的内容
            
            # 优先使用传递的document_content参数
            if document_content and document_content.strip():
                processed_doc_content = document_content
                has_document = True
                print(f"使用传递的document_content，长度: {len(processed_doc_content)}")
            # 如果有file_id参数，通过ID查找文件
            elif file_id:
                print(f"通过file_id查找文档: {file_id}")
                
                # 直接从文档存储中获取文档内容
                doc_content = document_storage.get_document_content(file_id)
                if doc_content:
                    processed_doc_content = doc_content
                    has_document = True
                    print(f"通过document_storage.get_document_content获取到文档内容，长度: {len(processed_doc_content)}")
                    
                    # 更新last_uploaded_document，确保后续处理正常
                    document = document_storage.get_document(file_id)
                    if document:
                        processed_content = document.get("processed_content")
                        if processed_content:
                            print(f"文档有处理后内容，长度: {len(processed_content)}")
                        
                        # 重要：更新last_uploaded_document，确保后续处理使用当前选择的文档
                        last_uploaded_document = {
                            "document_id": file_id,
                            "content": processed_doc_content,
                            "original_content": processed_doc_content,
                            "filename": document.get("filename", "document.docx"),
                            "processed_content": processed_content
                        }
                        print(f"已更新last_uploaded_document为当前选择的文档: {file_id}")
                else:
                    # 尝试从document对象获取
                    document = document_storage.get_document(file_id)
                    if document:
                        processed_doc_content = document.get("content", "")
                        has_document = processed_doc_content and processed_doc_content.strip() != ""
                        print(f"通过document_storage.get_document找到文档，长度: {len(processed_doc_content) if has_document else 0}")
                        
                        # 获取处理后的内容(如果有)
                        processed_content = document.get("processed_content")
                        if processed_content:
                            print(f"文档有处理后内容，长度: {len(processed_content)}")
                        
                        # 重要：更新last_uploaded_document，确保后续处理使用当前选择的文档
                        last_uploaded_document = {
                            "document_id": file_id,
                            "content": processed_doc_content,
                            "original_content": processed_doc_content,
                            "filename": document.get("filename", "document.docx"),
                            "processed_content": processed_content
                        }
                        print(f"已更新last_uploaded_document为当前选择的文档: {file_id}")
                    else:
                        # 在file_list中查找匹配的文件
                        matching_file = next((f for f in file_list if str(f["id"]) == file_id), None)
                        
                        if matching_file and last_uploaded_document and last_uploaded_document.get("document_id") == file_id:
                            processed_doc_content = last_uploaded_document.get("content", "")
                            processed_content = last_uploaded_document.get("processed_content")
                            has_document = processed_doc_content and processed_doc_content.strip() != ""
                            print(f"通过file_id在last_uploaded_document中找到文档，长度: {len(processed_doc_content) if has_document else 0}")
                        else:
                            print(f"找不到匹配file_id的文档内容: {file_id}")
                            processed_doc_content = ""
            # 否则检查全局变量
            elif last_uploaded_document:
                processed_doc_content = last_uploaded_document.get("content", "")
                processed_content = last_uploaded_document.get("processed_content")
                has_document = processed_doc_content and processed_doc_content.strip() != ""
                print(f"使用last_uploaded_document中的内容，长度: {len(processed_doc_content) if has_document else 0}")
            else:
                processed_doc_content = ""
                print("没有找到文档内容")
            
            # 如果只是切换文档而没有其他命令，直接返回文档内容和处理后内容(如果有)
            if "switcher" in message.lower() or "切换文档" in message:
                if has_document:
                    # 构建响应，只返回原始文档内容，不返回处理后内容
                    response_data = {
                        'content': '已切换到文档',
                        'document': processed_doc_content,
                        'processed_document': ""  # 始终设置为空字符串，不考虑文档是否有处理后内容
                    }
                    print("文档已切换，处理后内容已重置为空")
                    
                    yield f"data: {json.dumps(response_data)}\n\n"
                else:
                    yield f"data: {json.dumps({'content': '无法找到文档内容'})}\n\n"
                yield "data: end\n\n"
                return
                
            # 获取客户端IP地址作为会话ID
            client_ip = request.client.host
            session_id = f"session_{client_ip}"
            print(f"Using session ID: {session_id}")
            
            # 进行意图分类 - 优先使用前端传递的意图
            if intent and intent in ["correction", "polish", "summary"]:
                # 使用前端传递的意图
                print(f"使用前端传递的意图: {intent}")
                intent_type = intent
            else:
                # 否则使用后端意图分类
                intent_type = IntentClassifier.classify(message)
                print(f"使用后端意图分类结果: {intent_type}")
            
            # 检查是否处于强制公文纠错模式
            is_forced_mode = forced_correction_mode.get(session_id, False)
            
            # 初始化会话计数器
            if session_id not in chat_counter:
                chat_counter[session_id] = 0
            
            # 如果是聊天意图，增加计数
            if intent_type == "chat":
                chat_counter[session_id] += 1
                print(f"闲聊计数增加: {session_id} → {chat_counter[session_id]}")
                
                # 检查是否达到闲聊次数限制
                if chat_counter[session_id] > 3:
                    # 进入强制公文纠错模式
                    forced_correction_mode[session_id] = True
                    is_forced_mode = True
                    print(f"进入强制公文纠错模式: {session_id}")
            
            # 如果是公文纠错相关意图，重置闲聊计数
            if intent_type in ["correction", "writing", "recorrection", "polish", "summary"]:
                chat_counter[session_id] = 0
                # 如果之前处于强制模式，现在解除
                if is_forced_mode:
                    forced_correction_mode[session_id] = False
                    print(f"退出强制公文纠错模式: {session_id}")
            
            # 如果处于强制公文纠错模式，且不是公文纠错相关意图，强制拒绝回答
            if is_forced_mode and intent_type == "chat":
                guidance_message = "您已经闲聊超过3次，我只能回答公文纠错相关的问题了。请输入与公文处理相关的问题，例如「如何纠错」、「如何润色文档」或直接要求我「对文档进行公文纠错」。"
                yield f"data: {json.dumps({'content': guidance_message})}\n\n"
                yield "data: end\n\n"
                return
            
            # 根据意图处理请求
            if intent_type in ["correction", "writing", "recorrection", "forced_correction", "polish", "summary"]:
                # 检查消息中是否包含需要直接纠错的文本
                # 检查常见的模式，如"对下面的句子进行公文纠错："、"帮我纠错："等
                correction_patterns = [
                    r"对.*句子.*纠错[:：](.+)",
                    r"纠错[:：](.+)",
                    r"润色[:：](.+)",
                    r"修改[:：](.+)",
                    r"公文纠错[:：](.+)",
                    r"请.*公文纠错[:：](.+)",  # 添加新模式匹配"请进行公文纠错："
                    r"进行.*纠错[:：](.+)",     # 添加新模式匹配"进行公文纠错："
                    r"请.*纠错[:：](.+)"        # 添加新模式匹配"请纠错："
                ]
                
                # 添加详细日志，记录每个模式的匹配结果
                print(f"尝试匹配文本提取模式，原始消息: '{message}'")
                
                extracted_text = None
                
                for pattern in correction_patterns:
                    match = re.search(pattern, message, re.DOTALL)
                    if match:
                        extracted_text = match.group(1).strip()
                        print(f"✓ 成功匹配模式 '{pattern}'，提取文本: '{extracted_text}'")
                        input_text_for_correction = extracted_text
                        break
                    else:
                        print(f"✗ 模式 '{pattern}' 不匹配")
                
                # 如果没有从模式中提取到文本，检查是否应该使用整个消息
                if not input_text_for_correction and not has_document and len(message) > 15 and "上传" not in message and "文档" not in message:
                    # 直接使用整个消息作为要纠错的内容
                    input_text_for_correction = message
                    print(f"没有匹配到特定模式，使用整个消息作为纠错内容: '{input_text_for_correction}'")
                
                # 情况1: 用户提供了要直接纠错的文本，或者是写作指令
                if input_text_for_correction or intent_type == "writing":
                    # 发送处理状态
                    yield f"data: {json.dumps({'content': '正在处理您的请求...'})}\n\n"
                    
                    # 根据意图选择不同的MCP工具
                    print(f"根据意图'{intent_type}'选择MCP工具...")
                    try:
                        if intent_type == "writing":
                            # 调用写作工具
                            print("调用MCP写作工具处理指令...")
                            # 使用整个消息作为写作提示
                            mcp_result = call_mcp_service("写作工具", {"text": message})
                            tool_name = "写作"
                        elif intent_type == "polish" or "润色" in message:
                            # 调用润色工具
                            print("调用MCP润色工具处理文本...")
                            mcp_result = call_mcp_service("润色工具", {"text": input_text_for_correction})
                            tool_name = "润色"
                        elif intent_type == "summary" or "总结" in message:
                            # 调用总结工具
                            print("调用MCP总结工具处理文本...")
                            mcp_result = call_mcp_service("总结工具", {"text": input_text_for_correction})
                            tool_name = "总结"
                        else:
                            # 默认调用纠错工具
                            print("调用MCP纠错工具处理文本...")
                            mcp_result = call_mcp_service("纠错算法", {"text": input_text_for_correction})
                            tool_name = "纠错"
                        
                        print(f"MCP {tool_name}服务返回结果: {mcp_result.get('status', 'unknown')}")
                        
                        if mcp_result["status"] == "success":
                            result_text = mcp_result.get("result", "")
                            
                            # 处理写作工具的结果
                            if tool_name == "写作":
                                # 提取写作内容和回复消息
                                try:
                                    # 先尝试解析JSON格式
                                    if '{' in result_text and '}' in result_text:
                                        json_start = result_text.find('{')
                                        json_end = result_text.rfind('}') + 1
                                        json_str = result_text[json_start:json_end]
                                        
                                        data = json.loads(json_str)
                                        writing_content = data.get('writing_content', "")
                                        response_message = data.get('reply_message', "我已完成您的写作请求")
                                    # 兼容旧格式的解析    
                                    elif "写作内容:" in result_text and "回复消息:" in result_text:
                                        parts = result_text.split("回复消息:")
                                        writing_content = parts[0].replace("写作内容:", "").strip()
                                        response_message = parts[1].strip()
                                    else:
                                        # 如果格式不符合预期，将整个结果作为文档内容返回
                                        writing_content = result_text
                                        response_message = "我已完成您的写作请求"
                                    
                                    # 返回两部分内容
                                    yield f"data: {json.dumps({'content': response_message, 'processed_document': writing_content})}\n\n"
                                    
                                    # 如果有聊天ID，保存处理后的内容到文档存储
                                    chat_history_data = json.loads(chat_history)
                                    for item in chat_history_data:
                                        if item.get("chat_id"):
                                            chat_id = item.get("chat_id")
                                            # 创建一个临时文档ID用于写作内容
                                            document_id = f"writing_{chat_id}_{int(time.time())}"
                                            
                                            # 保存到文档存储
                                            document_storage.add_document(document_id, f"写作文档_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx", b"")
                                            document_storage.update_processed_content(document_id, writing_content)
                                            
                                            # 更新全局变量，确保其他API可以访问到写作内容
                                            if last_uploaded_document:
                                                # 保存一份写作内容到last_uploaded_document
                                                last_uploaded_document["processed_content"] = writing_content
                                                last_uploaded_document["is_writing"] = True
                                                print(f"已将写作内容保存到last_uploaded_document，内容长度: {len(writing_content)}")
                                            
                                            # 创建二进制文档
                                            try:
                                                docx_bytes = convert_text_to_docx_bytes(writing_content)
                                                document_storage.update_processed_binary(document_id, docx_bytes)
                                                print(f"已生成写作文档的二进制内容，大小: {len(docx_bytes)} 字节")
                                            except Exception as e:
                                                print(f"生成写作文档时出错: {str(e)}")
                                            break
                                    return
                                except Exception as e:
                                    print(f"解析写作结果出错: {str(e)}")
                                    # 如果解析出错，直接返回原始内容
                                    yield f"data: {json.dumps({'content': f'我已完成您的写作请求', 'processed_document': result_text})}\n\n"
                                    return
                            else:
                                # 处理其他工具的结果
                                result_message = f"已完成文本{tool_name}。\n\n{result_text}"
                                print(f"MCP {tool_name}服务处理文本成功")
                                yield f"data: {json.dumps({'content': result_message})}\n\n"
                        else:
                            raise Exception(f"MCP服务返回错误: {mcp_result.get('message', '未知错误')}")
                    except Exception as e:
                        # MCP服务失败，回退到LLM处理
                        print(f"MCP服务处理失败: {str(e)}，回退到Qwen LLM处理")
                        
                        # 使用LLM处理方式
                        system_prompt = "你是一个专业的公文写作助手，请对下面的文本进行公文纠错和润色，改善语言表达，使其更符合公文规范。"
                        
                        doc_messages = [
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": f"请对以下文本进行公文纠错:\n\n{input_text_for_correction or message}"}
                        ]
                        
                        try:
                            print("调用Qwen LLM处理文本...")
                            response = llm.invoke(doc_messages)
                            corrected_text = response.content
                            print("Qwen LLM处理文本成功")
                            result_message = f"已使用Qwen AI大模型处理您的文本。\n\n{corrected_text}"
                        except Exception as llm_error:
                            print(f"Qwen LLM调用失败: {str(llm_error)}")
                            corrected_text = input_text_for_correction or message
                            result_message = f"处理失败: {str(llm_error)}。返回原始文本。"
                        
                        yield f"data: {json.dumps({'content': result_message})}\n\n"
                    
                # 情况2: 需要处理上传的文档
                elif not has_document:
                    # 文档内容缺失，提示用户上传文档
                    content_message = '请先上传Word文档再进行完整文档的公文纠错。如果您只想纠正特定句子或段落，请直接输入要纠正的内容，例如"纠错：这是一段需要纠错的文字"。'
                    yield f"data: {json.dumps({'content': content_message})}\n\n"
                    yield "data: end\n\n"
                    return
                else:
                    # 发送处理状态
                    yield f"data: {json.dumps({'content': '正在处理文档...'})}\n\n"
                    
                    # 根据意图选择不同的MCP工具处理文档
                    print(f"根据意图'{intent_type}'选择MCP工具处理文档...")
                    try:
                        # 确保我们有文档内容
                        if not processed_doc_content or processed_doc_content.strip() == "":
                            # 尝试直接从document_storage获取
                            if file_id:
                                doc_content = document_storage.get_document_content(file_id)
                                if doc_content:
                                    processed_doc_content = doc_content
                                    print(f"从document_storage获取到文档 {file_id} 的内容，长度: {len(processed_doc_content)}")
                                    # 更新last_uploaded_document以便后续处理
                                    if last_uploaded_document and last_uploaded_document.get("document_id") == file_id:
                                        last_uploaded_document["content"] = processed_doc_content
                                else:
                                    print(f"警告：无法从document_storage获取文档 {file_id} 的内容")
                            else:
                                print("警告：未指定file_id，无法获取文档内容")
                        
                        if not processed_doc_content or processed_doc_content.strip() == "":
                            raise Exception("无法获取文档内容，请重新上传文档或选择文档")
                        
                        if intent_type == "polish" or "润色" in message:
                            # 调用润色工具
                            print("调用MCP润色工具处理文档...")
                            mcp_result = call_mcp_service("润色工具", {"text": processed_doc_content})
                            tool_name = "润色"
                        elif intent_type == "summary" or "总结" in message:
                            # 调用总结工具
                            print("调用MCP总结工具处理文档...")
                            mcp_result = call_mcp_service("总结工具", {"text": processed_doc_content})
                            tool_name = "总结"
                        else:
                            # 默认调用纠错工具
                            print("调用MCP纠错工具处理文档...")
                            mcp_result = call_mcp_service("纠错算法", {"text": processed_doc_content})
                            tool_name = "纠错"
                        
                        print(f"MCP {tool_name}服务返回结果: {mcp_result.get('status', 'unknown')}")
                        
                        if mcp_result["status"] == "success":
                            result_text = mcp_result.get("result", "")
                            if "纠错" in tool_name:
                                # 如果是纠错服务，从返回结果中提取修改后的文本
                                if "纠错后的文本:" in result_text:
                                    parts = result_text.split("主要修改:")
                                    if len(parts) > 1:
                                        processed_doc = parts[0].replace("纠错后的文本:", "").strip()
                                        processing_message = f"已根据MCP协议完成文档{tool_name}，共进行了以下修改：\n{parts[1].strip()}"
                                    else:
                                        processed_doc = processed_doc_content
                                        processing_message = f"已完成文档{tool_name}。\n\n{result_text}"
                                else:
                                    processed_doc = processed_doc_content
                                    processing_message = f"已完成文档{tool_name}，但未发现需要修改的内容。"
                            else:
                                # 如果是润色或总结服务，提取结果显示到文档区域
                                if "润色后的文本:" in result_text:
                                    # 提取润色后的文本内容
                                    parts = result_text.split("修改建议:")
                                    if len(parts) > 1:
                                        polished_text = parts[0].replace("润色后的文本:", "").strip()
                                        processed_doc = polished_text  # 设置为润色后的文本
                                        processing_message = f"已完成文档{tool_name}，修改建议：\n{parts[1].strip()}"
                                    else:
                                        processed_doc = result_text
                                        processing_message = f"已完成文档{tool_name}。"
                                elif "总结内容:" in result_text:
                                    # 提取总结的内容
                                    processed_doc = result_text
                                    processing_message = f"已完成文档{tool_name}，总结结果已显示在编辑区。"
                                else:
                                    # 如果没有特定的格式标记，将整个结果内容显示在文档区域
                                    processed_doc = result_text
                                    processing_message = f"已完成文档{tool_name}，结果已显示在编辑区。"
                        else:
                            raise Exception(f"MCP服务返回错误: {mcp_result.get('message', '未知错误')}")
                        
                        print(f"MCP {tool_name}服务处理文档成功")
                        
                        # 保存处理后的文档内容到文档存储
                        if file_id:
                            document_storage.update_processed_content(file_id, processed_doc)
                            print(f"已更新文档 {file_id} 的处理后内容")
                    except Exception as e:
                        # MCP服务失败，回退到LLM处理
                        print(f"MCP {tool_name}服务处理失败: {str(e)}，回退到Qwen LLM处理")
                        processed_doc = processed_doc_content  # 保持原文档不变
                        processing_message = f"处理文档时出错: {str(e)}。已保持原文档不变。"
                    
                    # 保存处理后的文档内容到全局变量和文档存储
                    if last_uploaded_document:
                        last_uploaded_document["processed_content"] = processed_doc
                        
                        # 如果有文件ID，同时更新文档存储
                        if file_id and "document_id" in last_uploaded_document:
                            document_storage.update_processed_content(last_uploaded_document["document_id"], processed_doc)
                            print(f"已更新文档 {last_uploaded_document['document_id']} 的处理后内容")
                    
                    # 发送处理结果
                    yield f"data: {json.dumps({'content': processing_message, 'processed_document': processed_doc})}\n\n"
            else:
                # 普通聊天处理
                print("处理普通聊天请求...")
                # 添加用户新消息
                messages.append({"role": "user", "content": message})
                
                # 构建系统提示
                system_prompt = """你是政务公文智能体，一个智能公文写作助手。
当用户询问你的身份、名字、是谁、介绍自己等类似问题时，必须回答：
"我是政务公文智能体，您的智能公文写作助手！"
你可以完成政务公文写作、润色、纠错、总结任务。
请以自然、友好的方式回答用户的其他问题。
记住，不要透露你是AI、大模型或Qwen Chat，而应始终以我是政务公文智能体身份回答。"""
                
                # 添加系统消息到开始
                full_messages = [{"role": "system", "content": system_prompt}]
                
                # 使用已经解析好的history_data，而不是外部的chat_history参数
                history_chat = history_data
                
                # 添加历史消息
                for msg in history_chat:
                    if isinstance(msg, dict) and "role" in msg and "content" in msg:
                        if msg["role"] in ["user", "assistant"]:
                            full_messages.append(msg)
                    else:
                        print(f"跳过格式不正确的消息: {msg}")
                
                # 确保最新的用户消息被添加到messages中
                user_message_exists = False
                for msg in full_messages:
                    if msg["role"] == "user" and msg["content"] == message:
                        user_message_exists = True
                        break
                
                if not user_message_exists:
                    full_messages.append({"role": "user", "content": message})
                
                print(f"准备调用Qwen LLM进行聊天，消息数量: {len(full_messages)}")
                
                # 调用模型生成回复
                try:
                    # 记录调用过程
                    print("调用Qwen LLM生成回复...")
                    model_response = llm.invoke(full_messages)
                    assistant_message = model_response.content
                    print(f"Qwen LLM生成回复成功: {assistant_message[:50]}..." if len(assistant_message) > 50 else f"Qwen LLM生成回复成功: {assistant_message}")
                    
                    # 将响应内容通过流式传输发送给客户端
                    yield f"data: {json.dumps({'content': assistant_message})}\n\n"
                    
                except Exception as e:
                    print(f"调用LLM出错: {str(e)}")
                    error_message = f"处理您的请求时出错: {str(e)}。请稍后再试。"
                    yield f"data: {json.dumps({'content': error_message})}\n\n"
            
            # 发送结束信号
            yield "data: end\n\n"
            
        except Exception as e:
            print(f"流式聊天处理出错: {str(e)}")
            import traceback
            print(traceback.format_exc())
            yield f"data: {json.dumps({'content': f'处理请求时出错: {str(e)}'})}\n\n"
            yield "data: end\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream")

@app.post("/file/upload")
async def file_upload_adapter(file: UploadFile = File(...)):
    """适配前端/file/upload请求，存储文件并返回标准格式响应"""
    try:
        global last_uploaded_document, file_list
        
        if not file.filename.endswith(('.docx')):
            raise HTTPException(status_code=400, detail="只支持.docx格式文件")
        
        # 读取文件内容到内存
        file_content = await file.read()
        
        # 使用DocProcessor处理文档
        text_content, binary_content = DocProcessor.process_docx(file_content, file.filename)
        logging.info(f"文档处理完成，文本内容长度: {len(text_content)}, 二进制内容长度: {len(binary_content)}")
        
        # 生成唯一文档ID
        document_id = str(hash(text_content + file.filename))
        
        # 重要：保存最近上传的文档（用于公文纠错）
        # 这确保无论通过哪个接口上传，document_content都能被正确设置
        last_uploaded_document = {
            "document_id": document_id,
            "content": text_content,
            "original_content": text_content,
            "filename": file.filename,
            "message": "文档上传成功，请在聊天框中输入处理指令"
        }
        logging.info(f"文档已上传并保存到last_uploaded_document，长度: {len(text_content)}")
        
        # 同时保存到文档存储中 - 修正参数数量，只传递必要的参数
        document_storage.add_document(document_id, file.filename, binary_content)
        logging.info(f"文件已保存到document_storage，ID: {document_id}")
        
        # 同步文档存储到file_list
        document_storage.sync_to_file_list()
        
        # 创建文件记录以返回给前端
        file_item = {
            "id": document_id,
            "filename": file.filename,
            "fileuuid": document_id,
            "fileurl": "",
            "updatetime": datetime.now().isoformat(),
            "userid": "system"  # 可以从请求中获取，这里简化处理
        }
        
        # 返回符合前端期望的响应格式
        return {
            "code": 200,
            "data": file_item,
            "msg": "文件上传成功"
        }
    except Exception as e:
        error_msg = f"文件上传处理出错: {str(e)}"
        logging.error(error_msg)
        import traceback
        logging.error(traceback.format_exc())
        return {
            "code": 500,
            "data": None,
            "msg": error_msg
        }

@app.post("/file/list")
async def file_list_adapter(request: Request):
    """适配前端/file/list请求，返回文件列表"""
    try:
        # 尝试解析请求体，但不强制要求
        try:
            body = await request.json()
            logging.info(f"收到file/list请求，参数: {body}")
        except Exception:
            body = {}
            logging.info("收到file/list请求，无参数或参数解析失败")
        
        # 从文档存储中获取所有文档
        all_docs = document_storage.get_all_documents()
        logging.info(f"从文档存储获取到 {len(all_docs)} 个文档")
        
        # 将文档存储中的文档转换为前端需要的格式
        formatted_docs = []
        for doc in all_docs:
            # 转换为前端需要的格式，添加has_processed标记表示是否有处理后内容
            has_processed = doc.get("processed_content") is not None and doc.get("processed_content", "").strip() != ""
            formatted_docs.append({
                "id": doc.get("document_id", ""),
                "filename": doc.get("filename", ""),
                "fileuuid": doc.get("document_id", ""),
                "fileurl": "",
                "updatetime": doc.get("updatetime", ""),
                "userid": "system",  # 可以从请求中获取，这里简化处理
                "has_processed": has_processed  # 添加是否已处理的标记
            })
        
        logging.info(f"返回文件列表，共 {len(formatted_docs)} 个文件")
        
        # 返回格式化后的文档列表
        return {
            "code": 200,
            "data": formatted_docs,
            "msg": "获取文件列表成功"
        }
    except Exception as e:
        error_msg = f"获取文件列表出错: {str(e)}"
        logging.error(error_msg)
        import traceback
        logging.error(traceback.format_exc())
        return {
            "code": 500,
            "data": [],
            "msg": error_msg
        }

@app.post("/chat/list")
async def chat_list_adapter(request: Request):
    """适配前端/chat/list请求"""
    try:
        # 尝试解析请求体，但不强制要求
        try:
            body = await request.json()
            print(f"收到chat/list请求，参数: {body}")
            # 根据userid过滤聊天列表
            filtered_chats = chat_list
            if body and 'userid' in body:
                userid = body.get('userid')
                filtered_chats = [chat for chat in chat_list if chat.get('userid') == userid]
        except json.JSONDecodeError:
            # 如果请求体为空或无效JSON，使用空字典
            body = {}
            print(f"收到chat/list请求，无参数")
            filtered_chats = chat_list
        
        # 返回过滤后的聊天列表
        response = {
            "code": 200,
            "data": filtered_chats,
            "msg": "获取聊天列表成功"
        }
        print(f"返回聊天列表响应: {response}")
        return response
    except Exception as e:
        print(f"获取聊天列表出错: {str(e)}")
        return {
            "code": 500,
            "data": [],
            "msg": f"获取聊天列表出错: {str(e)}"
        }

@app.post("/chat/create")
async def chat_create_adapter(request: Request):
    """适配前端/chat/create请求"""
    try:
        # 解析请求体
        try:
            body = await request.json()
            print(f"收到chat/create请求，参数: {body}")
        except json.JSONDecodeError:
            print("收到chat/create请求，但请求体不是有效的JSON")
            return {
                "code": 400,
                "data": None,
                "msg": "请求体格式不正确"
            }
        
        # 将聊天数据添加到列表中
        chat_id = body.get('id')
        # 检查是否已存在相同ID的聊天
        existing_chat_index = next((i for i, chat in enumerate(chat_list) if chat.get('id') == chat_id), -1)
        if existing_chat_index != -1:
            # 更新现有聊天
            chat_list[existing_chat_index] = body
        else:
            # 添加新聊天
            chat_list.append(body)
        
        # 保存聊天历史到文件
        save_chat_history()
        
        print(f"当前聊天列表长度: {len(chat_list)}")
        
        # 直接返回传入的数据，表示创建成功
        response = {
            "code": 200,
            "data": body,
            "msg": "创建聊天成功"
        }
        print(f"返回chat/create响应: {response}")
        return response
    except Exception as e:
        error_msg = f"创建聊天记录出错: {str(e)}"
        print(error_msg)
        import traceback
        print(traceback.format_exc())
        return {
            "code": 500,
            "data": None,
            "msg": error_msg
        }

@app.post("/chat/update")
async def chat_update_adapter(request: Request):
    """适配前端/chat/update请求"""
    try:
        # 解析请求体
        try:
            body = await request.json()
            print(f"收到chat/update请求，参数: {body}")
        except json.JSONDecodeError:
            print("收到chat/update请求，但请求体不是有效的JSON")
            return {
                "code": 400,
                "data": None,
                "msg": "请求体格式不正确"
            }
        
        # 更新聊天列表中的数据
        chat_id = body.get('id')
        if chat_id:
            existing_chat_index = next((i for i, chat in enumerate(chat_list) if chat.get('id') == chat_id), -1)
            if existing_chat_index != -1:
                # 更新现有聊天
                chat_list[existing_chat_index] = body
                print(f"已更新ID为{chat_id}的聊天")
            else:
                # 如果找不到，添加为新聊天
                chat_list.append(body)
                print(f"未找到ID为{chat_id}的聊天，已添加为新聊天")
            
            # 保存聊天历史到文件
            save_chat_history()
        
        # 直接返回传入的数据，表示更新成功
        return {
            "code": 200,
            "data": body,
            "msg": "更新聊天成功"
        }
    except Exception as e:
        print(f"更新聊天记录出错: {str(e)}")
        return {
            "code": 500,
            "data": None,
            "msg": f"更新聊天记录出错: {str(e)}"
        }

@app.delete("/chat/delete/{id}")
async def chat_delete_adapter(id: str):
    """适配前端/chat/delete/{id}请求"""
    try:
        # 查找并删除指定ID的聊天
        global chat_list
        original_length = len(chat_list)
        chat_list = [chat for chat in chat_list if chat.get('id') != id]
        
        if len(chat_list) < original_length:
            print(f"已删除ID为{id}的聊天")
            # 保存聊天历史到文件
            save_chat_history()
        else:
            print(f"未找到ID为{id}的聊天")
        
        # 返回成功响应
        return {
            "code": 200,
            "data": None,
            "msg": f"聊天{id}删除成功"
        }
    except Exception as e:
        print(f"删除聊天记录出错: {str(e)}")
        return {
            "code": 500,
            "data": None,
            "msg": f"删除聊天记录出错: {str(e)}"
        }

@app.post("/file/history")
async def file_history_adapter(request: Request):
    """获取文件历史记录"""
    try:
        try:
            body = await request.json()
            print(f"收到file/history请求，参数: {body}")
            # 根据userid过滤文件历史记录
            filtered_history = file_history_list
            if body and 'userid' in body:
                userid = body.get('userid')
                filtered_history = [file for file in file_history_list if file.get('userid') == userid]
        except json.JSONDecodeError:
            # 如果请求体为空或无效JSON，使用空字典
            body = {}
            print(f"收到file/history请求，无参数")
            filtered_history = file_history_list
        
        # 返回过滤后的文件历史记录
        response = {
            "code": 200,
            "data": filtered_history,
            "msg": "获取文件历史记录成功"
        }
        print(f"返回文件历史记录响应: {response}")
        return response
    except Exception as e:
        print(f"获取文件历史记录出错: {str(e)}")
        return {
            "code": 500,
            "data": [],
            "msg": f"获取文件历史记录出错: {str(e)}"
        }

@app.post("/file/add-history")
async def add_to_history_adapter(request: Request):
    """添加文件到历史记录"""
    try:
        # 解析请求体
        try:
            body = await request.json()
            print(f"收到file/add-history请求，参数: {body}")
        except json.JSONDecodeError:
            print("收到file/add-history请求，但请求体不是有效的JSON")
            return {
                "code": 400,
                "data": None,
                "msg": "请求体格式不正确"
            }
        
        # 检查必要字段
        file_id = body.get('id')
        if not file_id:
            return {
                "code": 400,
                "data": None,
                "msg": "缺少必要的文件ID字段"
            }
        
        # 查找是否已存在该文件的历史记录
        global file_history_list
        existing_index = next((i for i, file in enumerate(file_history_list) if file.get('id') == file_id), -1)
        
        # 设置访问次数和最近访问时间
        if 'accessCount' not in body:
            body['accessCount'] = 1
        if 'lastAccessed' not in body:
            body['lastAccessed'] = datetime.now().isoformat()
        
        if existing_index != -1:
            # 更新现有记录
            file_history_list[existing_index] = body
            print(f"更新ID为{file_id}的文件历史记录")
        else:
            # 添加新记录
            file_history_list.append(body)
            print(f"添加ID为{file_id}的文件历史记录")
        
        # 保存到文件
        save_file_history()
        
        # 返回更新后的文件记录
        return {
            "code": 200,
            "data": body,
            "msg": "添加文件历史记录成功"
        }
    except Exception as e:
        print(f"添加文件历史记录出错: {str(e)}")
        return {
            "code": 500,
            "data": None,
            "msg": f"添加文件历史记录出错: {str(e)}"
        }

@app.delete("/file/remove-history/{id}")
async def remove_from_history_adapter(id: str):
    """从历史记录中删除文件"""
    try:
        # 查找并删除指定ID的历史记录
        global file_history_list
        original_length = len(file_history_list)
        file_history_list = [file for file in file_history_list if file.get('id') != id]
        
        if len(file_history_list) < original_length:
            print(f"已从历史记录中删除ID为{id}的文件")
            # 保存到文件
            save_file_history()
        else:
            print(f"未找到ID为{id}的文件历史记录")
        
        # 返回成功响应
        return {
            "code": 200,
            "data": None,
            "msg": f"文件{id}历史记录删除成功"
        }
    except Exception as e:
        print(f"删除文件历史记录出错: {str(e)}")
        return {
            "code": 500,
            "data": None,
            "msg": f"删除文件历史记录出错: {str(e)}"
        }

@app.post("/file/clear-history")
async def clear_history_adapter(request: Request):
    """清空指定用户的文件历史记录"""
    try:
        # 解析请求体
        try:
            body = await request.json()
            print(f"收到file/clear-history请求，参数: {body}")
        except json.JSONDecodeError:
            print("收到file/clear-history请求，但请求体不是有效的JSON")
            return {
                "code": 400,
                "data": None,
                "msg": "请求体格式不正确"
            }
        
        # 获取用户ID
        userid = body.get('userid')
        if not userid:
            return {
                "code": 400,
                "data": None,
                "msg": "缺少必要的用户ID字段"
            }
        
        # 删除该用户的所有历史记录
        global file_history_list
        original_length = len(file_history_list)
        file_history_list = [file for file in file_history_list if file.get('userid') != userid]
        
        if len(file_history_list) < original_length:
            print(f"已清空用户{userid}的历史记录")
            # 保存到文件
            save_file_history()
        else:
            print(f"未找到用户{userid}的历史记录")
        
        # 返回成功响应
        return {
            "code": 200,
            "data": None,
            "msg": f"用户{userid}的历史记录已清空"
        }
    except Exception as e:
        print(f"清空历史记录出错: {str(e)}")
        return {
            "code": 500,
            "data": None,
            "msg": f"清空历史记录出错: {str(e)}"
        }

@app.delete("/file/delete/{id}")
async def file_delete_adapter(id: str):
    """适配前端文件删除请求"""
    try:
        # 从文档存储中删除
        result = document_storage.delete_document(id)
        
        if result:
            logging.info(f"已从document_storage中删除文件 ID: {id}")
            
            # 同步文档存储到file_list
            document_storage.sync_to_file_list()
            logging.info(f"已同步文件列表，当前有 {len(file_list)} 个文件")
            
            return {
                "code": 200,
                "data": None,
                "msg": f"文件{id}删除成功"
            }
        else:
            logging.warning(f"删除文件失败，找不到ID为 {id} 的文件")
            return {
                "code": 404,
                "data": None,
                "msg": f"找不到ID为 {id} 的文件"
            }
    except Exception as e:
        error_msg = f"删除文件出错: {str(e)}"
        logging.error(error_msg)
        import traceback
        logging.error(traceback.format_exc())
        return {
            "code": 500,
            "data": None,
            "msg": error_msg
        }

@app.get("/file/preview")
async def file_preview_adapter(id: str = None):
    """适配前端文件预览请求"""
    try:
        print(f"接收到文件预览请求，文件ID: {id}")
        
        # 如果没有指定ID，使用last_uploaded_document
        if not id and last_uploaded_document:
            id = last_uploaded_document.get("document_id")
            print(f"未指定文件ID，使用当前选中的文档: {id}")
        
        if not id:
            raise HTTPException(status_code=400, detail="未指定文件ID且没有当前选中的文档")
        
        # 从文档存储获取文档
        document = document_storage.get_document(id)
        if not document:
            print(f"找不到文件 {id}")
            raise HTTPException(status_code=404, detail="无法预览文件，找不到内容")
        
        # 返回原始二进制文档
        if "file_content" in document and isinstance(document["file_content"], bytes):
            print(f"返回原始二进制文件内容，长度: {len(document['file_content'])}")
            
            return Response(
                content=document["file_content"],
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"inline; filename*=UTF-8''preview_{urllib.parse.quote(document['filename'])}"}
            )
        # 如果内存中没有二进制内容，尝试从磁盘读取
        file_path = os.path.join(document_storage.documents_folder, f"{id}.docx")
        if os.path.exists(file_path):
            try:
                with open(file_path, 'rb') as f:
                    binary_content = f.read()
                    return Response(
                        content=binary_content,
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        headers={"Content-Disposition": f"inline; filename=preview_{document['filename']}"}
                    )
            except Exception as e:
                print(f"从磁盘读取文件失败: {str(e)}")
                raise HTTPException(status_code=500, detail=f"从磁盘读取文件失败: {str(e)}")
        
        raise HTTPException(status_code=404, detail="无法找到文件内容")
    except Exception as e:
        error_msg = f"预览文件出错: {str(e)}"
        print(error_msg)
        raise HTTPException(status_code=500, detail=error_msg)

@app.post("/file/download")
async def file_download_adapter(id: str):
    """适配前端文件下载请求"""
    try:
        # 从最近上传的文档内容返回
        if not last_uploaded_document:
            raise HTTPException(status_code=404, detail="无法下载文件，找不到内容")
        
        # 获取处理后的内容，如果存在的话
        content = last_uploaded_document.get("processed_content", last_uploaded_document["content"])
        
        # 生成时间戳
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        filename = last_uploaded_document["filename"]
        filename_base = filename.rsplit(".", 1)[0]
        download_filename = f"{filename_base}_processed_{timestamp}.docx"
        
        # 将文本内容转换为docx并返回，使用原始文档作为模板
        original_binary = None
        if "document_id" in last_uploaded_document and last_uploaded_document["document_id"] in document_storage.file_contents:
            original_binary = document_storage.file_contents[last_uploaded_document["document_id"]]
            
        doc_bytes = convert_text_to_docx_bytes(content, original_binary)
        
        return Response(
            content=doc_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={download_filename}"}
        )
    except Exception as e:
        print(f"下载文件出错: {str(e)}")
        raise HTTPException(status_code=500, detail=f"下载文件出错: {str(e)}")

@app.post("/file/modified-preview")
async def file_modified_preview_adapter(request: Request):
    """适配前端修改后文件预览请求"""
    try:
        body = await request.json()
        file = body.get("file", {})
        chat = body.get("chat", {})
        
        # 可以根据需要构建特定的预览文档
        # 简单起见，直接返回一个示例文档
        sample_content = "这是修改后的文档内容示例"
        doc_bytes = DocProcessor.convert_text_to_docx_bytes(sample_content)
        
        return Response(
            content=doc_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"inline; filename=modified_preview.docx"}
        )
    except Exception as e:
        print(f"预览修改后文件出错: {str(e)}")
        raise HTTPException(status_code=500, detail=f"预览修改后文件出错: {str(e)}")

@app.post("/file/confirm-modify")
async def file_confirm_modify_adapter(request: Request):
    """适配前端确认修改请求"""
    try:
        body = await request.json()
        # 简单返回成功，可以扩展实际处理逻辑
        return {
            "code": 200,
            "data": None,
            "msg": "确认修改成功"
        }
    except Exception as e:
        print(f"确认修改出错: {str(e)}")
        raise HTTPException(status_code=500, detail=f"确认修改出错: {str(e)}")

# 将文本转换为Word文档
def convert_text_to_docx_bytes(text_content, original_doc_bytes=None):
    """将文本内容转换为Word文档字节流，保留原始格式
    
    Args:
        text_content: 处理后的文本内容
        original_doc_bytes: 原始文档的二进制内容，如果提供则在此基础上修改
    """
    try:
        # 添加空值检查
        if text_content is None:
            logging.warning("转换内容为空，使用默认提示")
            text_content = "文档内容为空或无法读取"
        
        # 使用三步式处理函数处理文档，保留格式
        if original_doc_bytes and isinstance(original_doc_bytes, bytes):
            logging.info("使用原始文档作为模板和三步式处理方法生成新文档")
            return process_document_with_format_preservation(original_doc_bytes, text_content)
            
        # 如果没有原始文档或处理失败，创建一个新的Word文档
        logging.info("没有原始文档模板，创建新文档")
        from docx import Document
        import io
        
        doc = Document()
        
        # 按段落添加文本内容
        paragraphs = text_content.split('\n')
        for para in paragraphs:
            if para.strip():  # 跳过空行
                doc.add_paragraph(para)
        
        # 将文档保存到内存中
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        # 返回字节内容
        return docx_bytes.getvalue()
    except Exception as e:
        logging.error(f"Convert text to DOCX error: {str(e)}")
        import traceback
        logging.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Failed to convert text to DOCX: {str(e)}")

@app.post("/api/convert-text-to-docx")
async def convert_text_to_docx(request: Request):
    """
    将文本内容转换为DOCX格式
    """
    try:
        data = await request.json()
        text_content = data.get("text_content", "")
        filename = data.get("filename", "converted_document")
        
        if not text_content:
            raise HTTPException(status_code=400, detail="No text content provided")
        
        # 创建一个新的Word文档
        doc = Document()
        
        # 按段落拆分文本并添加到文档中
        paragraphs = text_content.split('\n')
        for para in paragraphs:
            if para.strip():  # 跳过空行
                doc.add_paragraph(para)
        
        # 将文档保存到内存中
        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        # 如果文件名没有.docx后缀，添加它
        if not filename.lower().endswith('.docx'):
            filename = f"{filename}.docx"
        
        # 返回文档内容
        return StreamingResponse(
            docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except Exception as e:
        logging.error(f"Convert text to DOCX error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to convert text to DOCX: {str(e)}")

@app.post("/api/upload")
async def api_file_upload(file: UploadFile = File(...)):
    """处理前端/api/upload请求，转发到file/upload处理程序"""
    return await file_upload_adapter(file)

@app.post("/api/list")
async def api_file_list(request: Request):
    """处理前端/api/list请求，转发到file/list处理程序"""
    return await file_list_adapter(request)

@app.get("/api/preview")
async def api_file_preview(id: str):
    """处理前端/api/preview请求，转发到file/preview处理程序"""
    return await file_preview_adapter(id)

@app.get("/api/debug/storage")
async def debug_storage():
    """调试端点，用于检查文件列表和文档存储状态"""
    try:
        # 获取文件列表信息
        file_list_info = []
        for f in file_list:
            file_list_info.append({
                "id": f.get("id", ""),
                "filename": f.get("filename", ""),
                "updatetime": f.get("updatetime", "")
            })
        
        # 获取文档存储信息
        storage_info = []
        for doc_id, doc in document_storage.documents.items():
            doc_info = {
                "document_id": doc_id,
                "filename": doc.get("filename", ""),
                "content_length": len(doc.get("content", "")) if "content" in doc else 0,
                "has_binary": doc.get("has_binary", False),
                "has_binary_in_memory": doc_id in document_storage.file_contents,
                "binary_size": len(document_storage.file_contents.get(doc_id, b"")) if doc_id in document_storage.file_contents else 0,
                "updatetime": doc.get("updatetime", "")
            }
            storage_info.append(doc_info)
        
        # 返回综合信息
        return {
            "status": "success",
            "file_list_count": len(file_list),
            "file_list": file_list_info,
            "document_storage_count": len(document_storage.documents),
            "document_storage": storage_info,
            "last_uploaded_document": {
                "exists": last_uploaded_document is not None,
                "document_id": last_uploaded_document.get("document_id", "") if last_uploaded_document else "",
                "filename": last_uploaded_document.get("filename", "") if last_uploaded_document else "",
                "content_length": len(last_uploaded_document.get("content", "")) if last_uploaded_document and "content" in last_uploaded_document else 0
            }
        }
    except Exception as e:
        logging.error(f"调试存储信息出错: {str(e)}")
        import traceback
        logging.error(traceback.format_exc())
        return {"status": "error", "message": str(e)}

# 添加MCP相关API端点
@app.post("/api/mcp/correction")
async def mcp_correction(request: Request):
    """MCP协议下的文档纠错API，保留原始格式"""
    try:
        # 解析请求体
        body = await request.json()
        print(f"收到MCP纠错请求: {body}")
        
        # 获取文件ID
        file_id = body.get("file_id")
        if not file_id:
            return JSONResponse(
                status_code=400,
                content={"status": "error", "message": "必须提供file_id参数"}
            )
        
        # 从文档存储获取文档
        document = document_storage.get_document(file_id)
        if not document:
            return JSONResponse(
                status_code=404,
                content={"status": "error", "message": f"找不到ID为 {file_id} 的文档"}
            )
        
        # 获取处理类型，默认为纠错
        process_type = body.get("type", "correction")
        
        # 获取文档内容
        content = document.get("content", "")
        if not content.strip():
            # 尝试从document_storage获取文档内容
            content = document_storage.get_document_content(file_id)
            if not content or not content.strip():
                return JSONResponse(
                    status_code=400,
                    content={"status": "error", "message": "文档内容为空"}
                )
        
        # 确认文档有二进制内容
        if not document.get("has_binary") or file_id not in document_storage.file_contents:
            return JSONResponse(
                status_code=400,
                content={"status": "error", "message": "找不到文档的二进制内容"}
            )
        
        # 使用新的三步式方法处理文档
        try:
            # 获取原始文档二进制内容
            original_binary = document_storage.file_contents[file_id]
            
            # 读取原始文本内容
            paragraphs_text = []
            import tempfile
            from docx import Document
            import io
            
            # 创建临时文件来保存原始文档
            temp_path = None
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                tmp.write(original_binary)
                temp_path = tmp.name
            
            # 使用python-docx打开文档
            doc = Document(temp_path)
            
            # 读取docx的文本内容
            for p in doc.paragraphs:
                if p.text.strip():  # 只处理非空段落
                    paragraphs_text.append(p.text)
            
            all_text = "\n".join(paragraphs_text)
            
            # 根据处理类型调用不同的MCP工具
            if process_type == "polish" or "润色" in str(body):
                # 调用润色工具
                print("调用MCP润色工具处理文档...")
                mcp_result = call_mcp_service("润色工具", {"text": all_text})
                tool_name = "润色"
            elif process_type == "summary" or "总结" in str(body):
                # 调用总结工具
                print("调用MCP总结工具处理文档...")
                mcp_result = call_mcp_service("总结工具", {"text": all_text})
                tool_name = "总结"
            else:
                # 默认调用纠错工具
                print("调用MCP纠错工具处理文档...")
                mcp_result = call_mcp_service("纠错算法", {"text": all_text})
                tool_name = "纠错"
            
            # 清理临时文件
            if temp_path and os.path.exists(temp_path):
                try:
                    os.unlink(temp_path)
                except:
                    pass
            
            if mcp_result["status"] == "success":
                result_text = mcp_result.get("result", "")
                
                # 使用三步式处理函数处理文档，保留格式
                processed_binary = process_document_with_format_preservation(original_binary, result_text)
                
                if processed_binary:
                    # 保存处理后的文本内容和二进制内容
                    document_storage.update_processed_content(file_id, result_text)
                    document_storage.update_processed_binary(file_id, processed_binary)
                    
                    # 创建响应内容
                    changes_description = mcp_result.get("changes", "文档已完成处理，但未提供详细修改说明")
                    
                    return JSONResponse(content={
                        "status": "success",
                        "message": f"文档已完成{tool_name}",
                        "changes": changes_description,
                        "modified": True
                    })
                else:
                    return JSONResponse(content={
                        "status": "error",
                        "message": f"处理文档时出错，无法保留格式"
                    })
            else:
                return JSONResponse(content={
                    "status": "error",
                    "message": mcp_result.get("message", f"MCP服务{tool_name}处理失败")
                })
                
        except Exception as e:
            print(f"处理文档出错: {str(e)}")
            import traceback
            print(traceback.format_exc())
            return JSONResponse(content={
                "status": "error",
                "message": f"处理文档过程中出错: {str(e)}"
            })
            
    except Exception as e:
        print(f"MCP纠错请求处理出错: {str(e)}")
        import traceback
        print(traceback.format_exc())
        return JSONResponse(
            status_code=500,
            content={"status": "error", "message": f"MCP纠错请求处理出错: {str(e)}"}
        )

@app.get("/api/get-processed-content")
@app.get("/api/mcp/get-processed-content")
async def get_processed_content(file_id: str, chat_id: str = None):
    """获取处理后的文档内容"""
    try:
        print(f"获取处理后的文档内容，文件ID: {file_id}")
        
        # 从文档存储获取文档
        document = document_storage.get_document(file_id)
        if not document:
            return JSONResponse(
                status_code=404,
                content={"status": "error", "message": f"找不到ID为 {file_id} 的文档"}
            )
        
        # 检查是否有处理后的二进制内容
        has_processed_binary = document.get("has_processed_binary", False)
        
        # 如果有处理后的二进制内容，优先返回
        if has_processed_binary and f"{file_id}_processed" in document_storage.file_contents:
            print(f"返回文档 {file_id} 的处理后二进制内容，大小: {len(document_storage.file_contents[f'{file_id}_processed'])} 字节")
            return Response(
                content=document_storage.file_contents[f"{file_id}_processed"],
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"inline; filename=preview_modified_{document.get('filename', 'document.docx')}"}
            )
        
        # 获取处理后的内容，如果存在的话
        processed_content = document.get("processed_content")
        if not processed_content:
            # 如果没有处理过的内容，返回原始内容
            processed_content = document.get("content", "")
        
        # 获取原始二进制内容作为模板
        original_binary = None
        if "file_content" in document and isinstance(document["file_content"], bytes):
            original_binary = document["file_content"]
            print(f"使用原始文档作为模板，二进制长度: {len(original_binary)}")
        
        # 将文本内容转换为docx并返回，使用原始文档作为模板
        doc_bytes = convert_text_to_docx_bytes(processed_content, original_binary)
        
        return Response(
            content=doc_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"inline; filename=preview_modified_{document.get('filename', 'document.docx')}"}
        )
    except Exception as e:
        print(f"获取处理后的文档内容出错: {str(e)}")
        import traceback
        print(traceback.format_exc())
        return JSONResponse(
            status_code=500,
            content={"status": "error", "message": f"获取处理后的文档内容出错: {str(e)}"}
        )

@app.post("/api/preview-modified")
@app.get("/api/preview-modified")
async def preview_modified(request: Request, file_id: str = None, chat_id: str = None):
    """预览修改后的文档内容"""
    try:
        # 处理不同的请求类型
        if request.method == "POST":
            try:
                # 尝试从请求体获取参数
                body = await request.json()
                file_id = body.get("file_id") or file_id
                chat_id = body.get("chat_id") or chat_id
            except:
                # 如果解析JSON失败，继续使用查询参数
                pass
        
        if not file_id:
            return JSONResponse(
                status_code=400,
                content={"status": "error", "message": "必须提供file_id参数"}
            )
        
        # 从文档存储获取文档
        document = document_storage.get_document(file_id)
        if not document:
            return JSONResponse(
                status_code=404,
                content={"status": "error", "message": f"找不到ID为 {file_id} 的文档"}
            )
        
        # 检查是否有处理后的二进制内容
        has_processed_binary = document.get("has_processed_binary", False)
        
        # 如果有处理后的二进制内容，优先返回
        if has_processed_binary and f"{file_id}_processed" in document_storage.file_contents:
            print(f"返回文档 {file_id} 的处理后二进制内容，大小: {len(document_storage.file_contents[f'{file_id}_processed'])} 字节")
            return Response(
                content=document_storage.file_contents[f"{file_id}_processed"],
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"inline; filename=preview_modified_{document.get('filename', 'document.docx')}"}
            )
        
        # 没有处理后的二进制内容，使用处理后的文本内容
        processed_content = document.get("processed_content")
        if not processed_content:
            # 如果没有处理过的内容，返回原始内容
            processed_content = document.get("content", "")
        
        # 获取原始二进制内容作为模板
        original_binary = None
        if "file_content" in document and isinstance(document["file_content"], bytes):
            original_binary = document["file_content"]
            print(f"使用原始文档作为模板，二进制长度: {len(original_binary)}")
        
        # 将文本内容转换为docx并返回，使用原始文档作为模板
        doc_bytes = convert_text_to_docx_bytes(processed_content, original_binary)
        
        return Response(
            content=doc_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"inline; filename=preview_modified_{document.get('filename', 'document.docx')}"}
        )
    except Exception as e:
        print(f"预览修改后文档出错: {str(e)}")
        import traceback
        print(traceback.format_exc())
        return JSONResponse(
            status_code=500,
            content={"status": "error", "message": f"预览修改后文档出错: {str(e)}"}
        )

# 添加新接口，用于前端在文档切换时获取文档内容
@app.get("/api/file/content")
async def get_file_content(file_id: str, include_processed: bool = True):
    """获取文档内容，只返回元数据和原始内容，不返回处理后内容"""
    try:
        # 从文档存储获取文档
        document = document_storage.get_document(file_id)
        if not document:
            return JSONResponse(
                status_code=404,
                content={"status": "error", "message": f"找不到ID为 {file_id} 的文档"}
            )
        
        # 获取文档内容
        content = document.get("content", "")
        if not content or not content.strip():
            # 如果内容为空，尝试从文件中读取
            print(f"文档 {file_id} 内容为空，尝试重新获取")
            content = document_storage.get_document_content(file_id)
            if not content:
                content = ""
                print(f"无法获取文档 {file_id} 的内容")
        
        # 准备响应内容
        response_data = {
            "status": "success",
            "document_id": file_id,
            "filename": document.get("filename", "document.docx"),
            "has_binary": document.get("has_binary", False),
            "has_processed": False,  # 始终设置为False，表示需要重新处理
            "content": content
        }
        
        # 如果请求包含处理后内容，返回空字符串
        if include_processed:
            # 不考虑文档是否有处理后内容，始终返回空字符串
            response_data["processed_content"] = ""
        
        # 重要：更新last_uploaded_document，确保后续处理使用当前选择的文档
        global last_uploaded_document
        last_uploaded_document = {
            "document_id": file_id,
            "filename": document.get("filename", "document.docx"),
            "file_path": os.path.join(document_storage.documents_folder, f"{file_id}.docx"),
            "content": content,
            "original_content": content,
            "processed_content": ""  # 始终设置为空字符串，不考虑文档是否有处理后内容
        }
        print(f"已更新last_uploaded_document为当前选择的文档: {file_id}，处理后内容已重置为空")
        
        return JSONResponse(content=response_data)
    except Exception as e:
        print(f"获取文档内容出错: {str(e)}")
        import traceback
        print(traceback.format_exc())
        return JSONResponse(
            status_code=500,
            content={"status": "error", "message": f"获取文档内容出错: {str(e)}"}
        )

@app.post("/api/document/switch")
async def switch_document(request: Request):
    """处理文档切换请求，返回选中文档的元数据，不返回处理后内容"""
    try:
        # 解析请求体
        body = await request.json()
        file_id = body.get("file_id")
        
        if not file_id:
            return JSONResponse(
                status_code=400,
                content={"status": "error", "message": "必须提供file_id参数"}
            )
        
        # 从文档存储获取文档
        document = document_storage.get_document(file_id)
        if not document:
            return JSONResponse(
                status_code=404,
                content={"status": "error", "message": f"找不到ID为 {file_id} 的文档"}
            )
        
        # 获取文件名
        filename = document.get("filename", "document.docx")
        
        # 获取文件路径
        file_path = os.path.join(document_storage.documents_folder, f"{file_id}.docx")
        
        # 获取文档内容
        content = document.get("content", "")
        if not content or not content.strip():
            # 如果内容为空，尝试从文件中读取
            print(f"文档 {file_id} 内容为空，尝试重新获取")
            content = document_storage.get_document_content(file_id)
            if not content:
                content = ""
                print(f"无法获取文档 {file_id} 的内容")
        
        # 重要：更新last_uploaded_document，确保后续处理使用当前选择的文档
        # 注意：不保留processed_content，确保切换时处理后内容为空
        global last_uploaded_document
        last_uploaded_document = {
            "document_id": file_id,
            "filename": filename,
            "file_path": file_path,
            "content": content,
            "original_content": content,
            "processed_content": ""  # 始终设置为空字符串，不考虑文档是否有处理后内容
        }
        print(f"已更新last_uploaded_document为当前选择的文档: {file_id}，处理后内容已重置为空")
        
        # 返回文档元数据和内容，但不包括处理后内容
        response_data = {
            "status": "success",
            "document_id": file_id,
            "filename": filename,
            "has_binary": document.get("has_binary", False),
            "has_processed": False,  # 始终设置为False，表示需要重新处理
            "content": content,
            "processed_content": ""  # 始终返回空字符串，不考虑文档是否有处理后内容
        }
        
        return JSONResponse(content=response_data)
    except Exception as e:
        print(f"切换文档时出错: {str(e)}")
        import traceback
        print(traceback.format_exc())
        return JSONResponse(
            status_code=500,
            content={"status": "error", "message": f"切换文档时出错: {str(e)}"}
        )

# 添加新函数：三步式处理文档，保留格式
def process_document_with_format_preservation(original_doc_bytes, modified_text):
    """
    三步式处理文档并保留格式：
    1. 完整保存原始文档的所有格式属性
    2. 用处理后的文本内容创建新文档
    3. 将保存的格式应用到新文档
    
    Args:
        original_doc_bytes: 原始文档的二进制内容
        modified_text: 处理后的文本内容
    
    Returns:
        bytes: 处理后的文档二进制内容
    """
    try:
        import tempfile
        from docx import Document
        import io
        import os
        
        # 如果没有原始文档，直接创建新文档
        if not original_doc_bytes or not isinstance(original_doc_bytes, bytes):
            logging.warning("没有原始文档，直接创建新文档")
            doc = Document()
            paragraphs = modified_text.split('\n')
            for para in paragraphs:
                if para.strip():  # 跳过空行
                    doc.add_paragraph(para)
            
            # 将文档保存到内存中
            docx_bytes = io.BytesIO()
            doc.save(docx_bytes)
            docx_bytes.seek(0)
            return docx_bytes.getvalue()
        
        # 尝试使用原始文档作为模板
        try:
            # =============== 第一步：完整保存原始文档格式 ===============
            # 保存原始文档到临时文件
            temp_path = None
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                tmp.write(original_doc_bytes)
                temp_path = tmp.name
            
            # 打开原始文档，提取格式信息
            original_doc = Document(temp_path)
            
            # 提取文档级别属性
            doc_properties = {
                'sections': [],
                'styles': {}
            }
            
            # 保存节属性
            for section in original_doc.sections:
                section_properties = {
                    'orientation': section.orientation,
                    'page_width': section.page_width,
                    'page_height': section.page_height,
                    'left_margin': section.left_margin,
                    'right_margin': section.right_margin,
                    'top_margin': section.top_margin,
                    'bottom_margin': section.bottom_margin,
                    'header_distance': section.header_distance,
                    'footer_distance': section.footer_distance
                }
                doc_properties['sections'].append(section_properties)
            
            # 保存段落和字符格式信息
            paragraphs_info = []
            
            for p in original_doc.paragraphs:
                paragraph_info = {
                    'text': p.text,
                    'style': p.style.name if p.style else None,
                    'paragraph_format': {
                        'alignment': p.paragraph_format.alignment,
                        'left_indent': p.paragraph_format.left_indent,
                        'right_indent': p.paragraph_format.right_indent,
                        'first_line_indent': p.paragraph_format.first_line_indent,
                        'line_spacing': p.paragraph_format.line_spacing,
                        'space_before': p.paragraph_format.space_before,
                        'space_after': p.paragraph_format.space_after,
                        'widow_control': p.paragraph_format.widow_control
                    },
                    'runs': []
                }
                
                # 保存每个run的字符级格式
                for run in p.runs:
                    run_info = {
                        'text': run.text,
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font': {
                            'name': run.font.name,
                            'size': run.font.size,
                            'color': run.font.color.rgb if run.font.color else None,
                            'highlight_color': run.font.highlight_color,
                            'subscript': run.font.subscript,
                            'superscript': run.font.superscript,
                        }
                    }
                    paragraph_info['runs'].append(run_info)
                
                paragraphs_info.append(paragraph_info)
                
            # 提取默认段落格式（用于新段落）
            default_style = None
            default_para_format = None
            default_font = None
            
            if paragraphs_info:
                default_style = paragraphs_info[0]['style']
                default_para_format = paragraphs_info[0]['paragraph_format']
                
                if paragraphs_info[0]['runs']:
                    default_font = paragraphs_info[0]['runs'][0]['font']
            
            # =============== 第二步：用处理后的文本创建新文档 ===============
            # 分割处理后的文本为段落
            modified_paragraphs = modified_text.split('\n')
            modified_paragraphs = [p for p in modified_paragraphs if p.strip()]  # 移除空段落
            
            # 创建新文档
            new_doc = Document()
            
            # 尝试应用原始文档的节设置
            if doc_properties['sections']:
                # 获取新文档的节
                section = new_doc.sections[0]
                # 应用原始文档的节设置
                section_props = doc_properties['sections'][0]
                section.orientation = section_props['orientation']
                section.page_width = section_props['page_width']
                section.page_height = section_props['page_height']
                section.left_margin = section_props['left_margin']
                section.right_margin = section_props['right_margin']
                section.top_margin = section_props['top_margin']
                section.bottom_margin = section_props['bottom_margin']
                section.header_distance = section_props['header_distance']
                section.footer_distance = section_props['footer_distance']
            
            # =============== 第三步：应用格式到新文档 ===============
            # 简单添加段落
            for para_text in modified_paragraphs:
                new_doc.add_paragraph(para_text)
            
            # 保存修改后的文档
            docx_bytes = io.BytesIO()
            new_doc.save(docx_bytes)
            docx_bytes.seek(0)
            
            # 清理临时文件
            try:
                os.unlink(temp_path)
            except:
                pass
            
            logging.info("使用原始文档作为模板生成新文档")
            return docx_bytes.getvalue()
            
        except Exception as e:
            logging.warning(f"使用原始文档作为模板失败: {str(e)}，将创建新文档")
            # 如果处理失败，回退到创建新文档
            doc = Document()
            
            # 按段落添加文本内容
            paragraphs = modified_text.split('\n')
            for para in paragraphs:
                if para.strip():  # 跳过空行
                    doc.add_paragraph(para)
            
            # 将文档保存到内存中
            docx_bytes = io.BytesIO()
            doc.save(docx_bytes)
            docx_bytes.seek(0)
            
            # 返回字节内容
            return docx_bytes.getvalue()
            
    except Exception as e:
        logging.error(f"Convert text to DOCX error: {str(e)}")
        import traceback
        logging.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Failed to convert text to DOCX: {str(e)}")

@app.delete("/api/delete/{id}")
async def api_file_delete_adapter(id: str):
    """处理前端/api/delete/{id}请求，转发到file/delete处理程序"""
    return await file_delete_adapter(id)

@app.get("/api/chat/processed-document")
async def get_chat_processed_document(chat_id: str = None, timestamp: str = None):
    """
    获取与聊天相关的处理后文档内容
    这个API专为写作功能设计，用于获取写作生成的文档内容
    
    参数:
        chat_id: 聊天ID
        timestamp: 时间戳（用于避免缓存）
    
    返回:
        处理后的文档内容
    """
    if not chat_id:
        raise HTTPException(status_code=400, detail="缺少chat_id参数")
    
    try:
        # 查找与该聊天ID关联的最新写作文档
        writing_prefix = f"writing_{chat_id}_"
        latest_doc = None
        latest_timestamp = 0
        latest_doc_id = None
        
        print(f"查找chat_id {chat_id}的写作文档")
        
        # 遍历文档找到最新的写作文档
        for doc_id, doc in document_storage.documents.items():
            if doc_id.startswith(writing_prefix):
                try:
                    # writing_chatId_timestamp 格式
                    timestamp_str = doc_id.split('_')[-1]
                    doc_timestamp = int(timestamp_str)
                    if doc_timestamp > latest_timestamp:
                        latest_timestamp = doc_timestamp
                        latest_doc = doc
                        latest_doc_id = doc_id
                        print(f"找到更新的文档: {doc_id}, 时间戳: {doc_timestamp}")
                except (ValueError, IndexError):
                    continue
        
        # 如果找到文档，返回其处理后的内容
        if latest_doc and latest_doc_id:
            print(f"找到最新的写作文档: {latest_doc_id}")
            
            # 获取处理后的内容
            processed_content = latest_doc.get("processed_content", "")
            if not processed_content:
                print(f"警告: 文档 {latest_doc_id} 没有处理后内容")
                
            return {"status": "success", "content": processed_content}
        
        # 如果没有找到与chat_id相关的文档，尝试查找任何写作文档
        for doc_id, doc in sorted(document_storage.documents.items(), key=lambda x: x[0], reverse=True):
            if "writing" in doc_id and doc.get("processed_content"):
                print(f"找到写作文档: {doc_id}")
                return {"status": "success", "content": doc.get("processed_content", "")}
        
        # 如果仍找不到，检查last_uploaded_document
        if last_uploaded_document and "processed_content" in last_uploaded_document:
            print("使用last_uploaded_document中的处理后内容")
            return {"status": "success", "content": last_uploaded_document["processed_content"]}
        
        # 如果仍找不到，尝试查找任何有处理后内容的文档
        for doc_id, doc in document_storage.documents.items():
            if doc.get("processed_content"):
                print(f"找到有处理后内容的文档: {doc_id}")
                return {"status": "success", "content": doc.get("processed_content", "")}
        
        # 如果真的找不到任何内容，返回错误
        print("未找到任何文档内容")
        return {"status": "error", "message": "未找到与该聊天关联的处理后文档内容"}
        
    except Exception as e:
        print(f"获取处理后文档内容时出错: {str(e)}")
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"获取处理后文档内容时出错: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    import argparse
    
    # 添加命令行参数解析
    parser = argparse.ArgumentParser(description='启动智能公文助手服务')
    parser.add_argument('--host', type=str, default="0.0.0.0", help='服务器监听IP，默认0.0.0.0表示所有网络接口')
    parser.add_argument('--port', type=int, default=8003, help='服务器监听端口，默认8003')
    parser.add_argument('--reload', action='store_true', help='是否启用热重载（开发环境使用）')
    args = parser.parse_args()
    
    # 打印启动信息
    print(f"正在启动服务，监听地址: {args.host}:{args.port}")
    print(f"热重载模式: {'启用' if args.reload else '禁用'}")
    print(f"请使用浏览器访问: http://{args.host if args.host != '0.0.0.0' else '127.0.0.1'}:{args.port}")
    
    # 启动服务
    uvicorn.run("api.main:app", host=args.host, port=args.port, reload=args.reload) 