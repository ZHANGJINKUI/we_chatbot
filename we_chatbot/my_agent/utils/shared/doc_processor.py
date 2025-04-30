from langchain_community.document_loaders import Docx2txtLoader
from docx import Document  # 正确导入 python-docx
import io
import os
from typing import Optional, Tuple
import logging



class DocProcessor:
    @staticmethod
    def load_doc(file_path):
        """优化的文档加载方法，直接使用python-docx提高效率"""
        try:
            # 直接使用python-docx读取文档
            doc = Document(file_path)
            
            # 使用列表推导式快速提取段落内容
            content = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
            
            # 改进错误处理
            if not content:
                print("警告：文档内容为空，返回默认值")
                return "请提供有效的文档内容"
            
            return content
            
        except Exception as e:
            print(f"加载文档时出现错误: {e}")
            return "加载文档时出现错误，请检查文档格式"
    
    @staticmethod
    def load_doc_stream(file_stream, filename: Optional[str] = None):
        """从流中直接读取文档，避免保存临时文件的开销"""
        try:
            # 记录日志
            logging.info(f"开始从流中读取文档，文件名: {filename}, 数据类型: {type(file_stream)}, 大小: {len(file_stream)}")
            
            # 确保文件流是字节类型
            if not isinstance(file_stream, bytes):
                raise ValueError(f"文件流类型错误，预期bytes但收到: {type(file_stream)}")
            
            # 直接从内存流读取
            doc = Document(io.BytesIO(file_stream))
            
            # 使用列表推导式快速提取段落内容
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # 检查是否有段落内容
            if not paragraphs:
                logging.warning("文档中没有找到文本段落，将检查表格内容")
                
                # 尝试从表格中提取内容
                table_texts = []
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ' | '.join(cell.text for cell in row.cells if cell.text.strip())
                        if row_text:
                            table_texts.append(row_text)
                
                if table_texts:
                    content = '\n'.join(table_texts)
                    logging.info(f"从表格中提取了内容，总长度: {len(content)}")
                else:
                    # 如果表格中也没有内容，返回默认值
                    logging.warning("文档内容为空，返回默认内容")
                    return "请提供有效的文档内容，当前文档内容为空"
            else:
                # 正常情况下合并段落
                content = '\n'.join(paragraphs)
                logging.info(f"成功从文档中提取了{len(paragraphs)}个段落，总长度: {len(content)}")
            
            return content
            
        except Exception as e:
            error_msg = f"从流加载文档时出现错误: {str(e)}"
            logging.error(error_msg)
            return error_msg
    
    @staticmethod
    def save_doc(content: str, output_path: str):
        """保存结果到新文档"""
        try:
            doc = Document()  # 直接使用 Document
            
            # 按段落分割并添加到文档中
            for paragraph in content.split('\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph)
            
            # 确保目录存在
            os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
            
            # 保存文档
            doc.save(output_path)
            logging.info(f"文档已保存到: {output_path}")
            return True
        except Exception as e:
            error_msg = f"保存文档出错: {str(e)}"
            logging.error(error_msg)
            return False
    
    @staticmethod
    def save_text_to_docx(text_content: str, output_path: str):
        """将文本内容保存为Word文档"""
        try:
            # 创建一个新的Word文档
            doc = Document()
            
            # 检查文本内容是否为空
            if text_content is None:
                text_content = "文档内容为空"
            
            # 按段落拆分文本并添加到文档中
            paragraphs = text_content.split('\n')
            for para in paragraphs:
                if para.strip():  # 跳过空行
                    doc.add_paragraph(para)
            
            # 确保目录存在
            os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
                    
            # 保存文档
            doc.save(output_path)
            logging.info(f"文本已转换为Word文档并保存到: {output_path}")
            return True
        except Exception as e:
            error_msg = f"保存文本为Word文档出错: {str(e)}"
            logging.error(error_msg)
            return False
    
    @staticmethod
    def process_docx(file_content: bytes, filename: str) -> Tuple[str, bytes]:
        """处理Word文档，返回文本内容和原始二进制内容"""
        try:
            # 从二进制数据读取文档内容
            text_content = DocProcessor.load_doc_stream(file_content, filename)
            
            # 记录日志
            logging.info(f"成功处理Word文档: {filename}, 提取文本长度: {len(text_content)}")
            
            return text_content, file_content
        except Exception as e:
            error_msg = f"处理Word文档时出错: {str(e)}"
            logging.error(error_msg)
            # 返回错误信息作为文本内容
            return error_msg, file_content

    @staticmethod
    def process_docx_file(file_path: str) -> Tuple[str, bytes]:
        """从文件路径处理Word文档，返回文本内容和原始二进制内容"""
        try:
            # 读取文件的二进制内容
            with open(file_path, 'rb') as f:
                file_content = f.read()
            
            # 从文件加载文档内容
            text_content = DocProcessor.load_doc(file_path)
            
            # 记录日志
            filename = os.path.basename(file_path)
            logging.info(f"成功处理Word文档文件: {filename}, 提取文本长度: {len(text_content)}")
            
            return text_content, file_content
        except Exception as e:
            error_msg = f"处理Word文档文件时出错: {str(e)}"
            logging.error(error_msg)
            # 返回错误信息作为文本内容
            return error_msg, b""