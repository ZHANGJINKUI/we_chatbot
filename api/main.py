from fastapi import FastAPI, UploadFile, File, HTTPException, Body, Request, Response, Query, Header, BackgroundTasks, Depends
from fastapi.responses import FileResponse, StreamingResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
import os
import tempfile
import sys
import shutil
import json
import logging
from typing import List, Dict, Any, Optional
from pydantic import BaseModel
# 注释掉数据库导入
# from admin.auth.routes import router as auth_router
# from admin.database import Base, engine
# 导入DocProcessor和相关组件
from my_agent.utils.shared.doc_processor import DocProcessor
from my_agent.utils.shared.intent_classifier import IntentClassifier
from langchain_openai import ChatOpenAI
import os
from dotenv import load_dotenv
import re
import asyncio
from io import BytesIO
from datetime import datetime
from pathlib import Path
# 从my_agent.agent_tools导入CSC和调用MCP的方法
from my_agent.agent_tools import csc as mcp_csc, call_mcp_service
# 直接导入完整的模块以便于调试
import my_agent.agent_tools
from docx import Document

# 定义存储文件路径
DATA_DIR = Path(__file__).parent / "data"
CHAT_HISTORY_FILE = DATA_DIR / "chat_history.json"
FILE_LIST_FILE = DATA_DIR / "file_list.json"

# 确保数据目录存在
DATA_DIR.mkdir(parents=True, exist_ok=True)

# 全局变量，用于存储最近上传的文档
last_uploaded_document = None
# 全局变量，用于存储当前文件列表
file_list = []
# 全局变量，记录用户会话是否进入强制公文纠错模式
forced_correction_mode = {}
# 全局变量，存储聊天历史记录
chat_list = []
# 全局变量存储处理后的文档
last_uploaded_document = None
# 全局变量存储聊天历史
chat_list = []
# 全局变量存储文件历史记录
file_history_list = []

# 保存聊天历史的文件路径
CHAT_HISTORY_FILE = "chat_history.json"
# 保存文件历史记录的文件路径
FILE_HISTORY_FILE = "file_history.json"

# 加载聊天历史
def load_chat_history():
    global chat_list
    try:
        if os.path.exists(CHAT_HISTORY_FILE):
            with open(CHAT_HISTORY_FILE, "r", encoding="utf-8") as f:
                chat_list = json.load(f)
                print(f"从文件加载聊天历史，共 {len(chat_list)} 条记录")
    except Exception as e:
        print(f"加载聊天历史出错: {str(e)}")
        chat_list = []

# 保存聊天历史到文件
def save_chat_history():
    try:
        with open(CHAT_HISTORY_FILE, "w", encoding="utf-8") as f:
            json.dump(chat_list, f, ensure_ascii=False, indent=2)
            print(f"聊天历史已保存到文件，共 {len(chat_list)} 条记录")
    except Exception as e:
        print(f"保存聊天历史出错: {str(e)}")

# 加载文件历史记录
def load_file_history():
    global file_history_list
    try:
        if os.path.exists(FILE_HISTORY_FILE):
            with open(FILE_HISTORY_FILE, "r", encoding="utf-8") as f:
                file_history_list = json.load(f)
                print(f"从文件加载文件历史记录，共 {len(file_history_list)} 条记录")
    except Exception as e:
        print(f"加载文件历史记录出错: {str(e)}")
        file_history_list = []

# 保存文件历史记录到文件
def save_file_history():
    try:
        with open(FILE_HISTORY_FILE, "w", encoding="utf-8") as f:
            json.dump(file_history_list, f, ensure_ascii=False, indent=2)
            print(f"文件历史记录已保存到文件，共 {len(file_history_list)} 条记录")
    except Exception as e:
        print(f"保存文件历史记录出错: {str(e)}")

# 首先创建FastAPI应用实例
app = FastAPI(title="文档处理API")

# 在应用启动时加载聊天历史
@app.on_event("startup")
async def startup_event():
    # 初始化时加载聊天历史
    load_chat_history()
    # 初始化时加载文件历史记录
    load_file_history()
    print("应用启动，已加载聊天历史和文件历史记录")

# 配置CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # 允许所有源
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 定义临时文件路径
TEMP_DIR = tempfile.gettempdir()

# 加载环境变量和初始化模型
# 明确指定.env文件路径
dotenv_path = '/home/zjk/a/GW/we_chatbot/.env'
load_dotenv(dotenv_path)
deepseek_api_key = os.getenv("DEEPSEEK_API_KEY")
deepseek_base_url = os.getenv("DEEPSEEK_API_BASE")

# 检查环境变量是否设置
if not deepseek_api_key or not deepseek_base_url:
    print(f"警告: DeepSeek API 配置缺失。尝试从 {dotenv_path} 加载但未成功。")
    print("将使用默认配置继续...")
else:
    print(f"成功从 {dotenv_path} 加载 DeepSeek API 配置")

# 设置环境变量
os.environ["OPENAI_API_KEY"] = deepseek_api_key or "dummy_key"
os.environ["OPENAI_API_BASE"] = deepseek_base_url or "https://api.deepseek.com/v1"

# 初始化LLM，使用正确的模型名称
try:
    llm = ChatOpenAI(
        model="deepseek-chat", 
        temperature=0.7,
        api_key=deepseek_api_key,
        base_url=deepseek_base_url
    )
    print("成功初始化DeepSeek LLM模型")
except Exception as e:
    print(f"初始化DeepSeek LLM模型失败: {str(e)}")
    print("将使用备用配置...")
    # 备用配置，使用最基本的设置
    llm = ChatOpenAI(model="deepseek-chat")

# 闲聊计数器 - 全局变量
chat_counter = {}

# 定义请求和响应模型
class UploadDocumentResponse(BaseModel):
    document_id: str
    content: str
    message: str

class ChatRequest(BaseModel):
    message: str
    document_content: str = None  # 新增文档内容字段
    chat_history: list = []

class ChatResponse(BaseModel):
    response: str
    chat_history: list
    processed_document: str = None  # 新增处理后的文档内容字段

class SaveDocumentRequest(BaseModel):
    document_id: str
    content: str
    filename: str = "processed_document.docx"

# 添加登录请求和响应模型
class LoginRequest(BaseModel):
    userid: str
    password: str

class LoginResponseData(BaseModel):
    user: dict
    token: str

class LoginResponse(BaseModel):
    code: int
    data: LoginResponseData
    msg: str

# 定义路由
@app.get("/")
def read_root():
    return {"message": "Welcome to the User Management System"}

# 添加健康检查API
@app.get("/api/health-check")
async def health_check():
    """健康检查端点"""
    return {"status": "ok", "message": "Service is running"}

# 添加简单的登录API端点
@app.post("/api/login", response_model=LoginResponse)
async def login(login_request: LoginRequest):
    """处理用户登录请求"""
    # 这里简化处理，只要提供了用户名和密码就认为登录成功
    if login_request.userid and login_request.password:
        # 返回模拟的用户信息和token，包装为前端期望的格式
        response_data = {
            "user": {
                "userid": login_request.userid,
                "username": login_request.userid,
                "role": "user"
            },
            "token": "mock_token_123456789"
        }
        
        # 包装为前端期望的格式
        return {
            "code": 200,
            "data": response_data,
            "msg": "登录成功"
        }
    else:
        raise HTTPException(status_code=401, detail="用户名或密码错误")

# 修改文档上传函数，添加对file为None的处理
@app.post("/api/upload-document")
async def upload_document(file: UploadFile = None):
    """上传Word文档并返回内容，不进行处理"""
    # 存储最近上传的文档内容
    global last_uploaded_document
    
    # 如果file为None，返回上次上传的文档内容
    if file is None:
        if not last_uploaded_document:
            raise HTTPException(status_code=404, detail="没有已上传的文档")
        return last_uploaded_document
    
    if not file.filename.endswith(('.docx')):
        raise HTTPException(status_code=400, detail="只支持.docx格式文件")
    
    try:
        # 读取文件内容到内存
        file_content = await file.read()
        
        # 直接从内存流读取文档，避免临时文件I/O开销
        content = DocProcessor.load_doc_stream(file_content, file.filename)
        
        # 生成唯一文档ID
        document_id = str(hash(content + file.filename))
        
        # 保存最近上传的文档
        last_uploaded_document = {
            "document_id": document_id,
            "content": content,
            "original_content": content,
            "filename": file.filename,
            "message": "文档上传成功，请在聊天框中输入处理指令"
        }
        
        return last_uploaded_document
        
    except Exception as e:
        print(f"处理文档时出错: {str(e)}")
        raise HTTPException(status_code=500, detail=f"处理文档时出错: {str(e)}")

@app.post("/api/process-document", deprecated=True)
async def process_document_api(file: UploadFile = File(...)):
    """上传并处理Word文档 (已弃用，保留兼容)"""
    return await upload_document(file)

# 修改聊天接口，处理带有文档内容的请求
@app.post("/api/chat", response_model=ChatResponse)
async def chat_endpoint(chat_request: ChatRequest, request: Request):
    """处理聊天请求，支持文档处理"""
    try:
        print(f"Received chat request: {chat_request.message[:100]}...")
        
        # 准备消息历史
        messages = chat_request.chat_history.copy()
        
        # 检查是否有文档内容，如果没有，尝试从最近上传的文档中获取
        document_content = chat_request.document_content
        if (document_content is None or document_content.strip() == "") and last_uploaded_document:
            document_content = last_uploaded_document["content"]
            print(f"使用最近上传的文档内容，长度: {len(document_content)}")
        
        has_document = document_content is not None and document_content.strip() != ""
        
        # 进行意图分类
        intent = IntentClassifier.classify(chat_request.message)
        print(f"Intent classified as: {intent}")
        
        # 获取客户端IP地址作为会话ID
        client_ip = request.client.host
        session_id = f"session_{client_ip}"
        print(f"Using session ID: {session_id}")
        
        # 初始化会话计数器
        if session_id not in chat_counter:
            chat_counter[session_id] = 0
            print(f"Initialized chat counter for {session_id}")
            
        # 如果是聊天意图，增加计数
        if intent == "chat":
            chat_counter[session_id] += 1
            print(f"Increased chat counter for {session_id} to {chat_counter[session_id]}")
            
        # 检查是否超过闲聊次数限制
        if intent == "chat" and chat_counter[session_id] > 3:
            # 重置意图为强制文档处理
            intent = "forced_correction"
            # 添加用户消息到历史
            messages.append({"role": "user", "content": chat_request.message})
            # 添加系统消息，引导用户回到公文纠错
            guidance_message = "您已经闲聊超过3次，请回到公文纠错相关任务。您可以尝试输入「纠错」、「润色」等指令来处理您的文档。"
            messages.append({"role": "assistant", "content": guidance_message})
            
            return {
                "response": guidance_message,
                "chat_history": messages,
                "processed_document": None
            }
            
        # 根据意图处理
        if intent in ["correction", "writing", "recorrection", "forced_correction"]:
            # 检查是否有文档内容
            if not has_document:
                # 文档内容缺失，提示用户上传文档
                messages.append({"role": "user", "content": chat_request.message})
                no_doc_message = "请先上传Word文档再进行公文纠错。您可以点击左侧面板的「上传文档」按钮上传文件。"
                messages.append({"role": "assistant", "content": no_doc_message})
                
                return {
                    "response": no_doc_message,
                    "chat_history": messages,
                    "processed_document": None
                }
            
            # 修改为使用MCP协议封装的CSC纠错服务
            print("使用MCP协议的CSC服务处理文档...")
            
            # 调用MCP服务进行文档纠错
            csc_result = mcp_csc(document_content)
            
            if csc_result["status"] == "success":
                # 获取处理后的文档内容
                if csc_result.get("modified", False):
                    processed_document = csc_result["corrected_text"]
                    processing_message = f"文档纠错完成，共进行了以下修改：\n{csc_result.get('changes', '未提供详细修改说明')}"
                else:
                    processed_document = document_content
                    processing_message = csc_result.get("message", "文档未发现需要纠错的内容，已保持原文。")
                
                # 将用户消息添加到聊天历史
                messages.append({"role": "user", "content": chat_request.message})
                
                # 将处理结果添加到聊天历史
                success_message = f"已完成文档纠错。{processing_message}"
                messages.append({"role": "assistant", "content": success_message})
                
                # 保存处理后的文档到全局变量
                if last_uploaded_document:
                    last_uploaded_document["processed_content"] = processed_document
                
                # 返回处理结果和更新后的聊天历史
                return {
                    "response": success_message,
                    "chat_history": messages,
                    "processed_document": processed_document
                }
            else:
                # 处理失败，回退到原始的LLM处理方式
                print("MCP服务处理失败，回退到LLM处理...")
                error_message = csc_result.get("message", "MCP服务处理失败")
                print(f"错误: {error_message}")
                
            # 构建包含文档内容的系统提示
            system_prompt = "你是一个专业的文档处理助手。请根据用户的指令处理以下文档内容。"
            
            # 将文档内容和用户指令组合成完整提示
            full_prompt = f"文档内容:\n{document_content}\n\n用户指令:\n{chat_request.message}"
            
            # 准备消息序列
            doc_messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": full_prompt}
            ]
            
            # 调用模型处理文档
            print("转为调用大语言模型处理文档...")
            response = llm.invoke(doc_messages)
            
            # 将用户消息添加到聊天历史
            messages.append({"role": "user", "content": chat_request.message})
            
            # 将模型响应添加到聊天历史
            fallback_message = "MCP服务暂时不可用，已使用备用方法处理您的文档。处理结果已显示在中间区域。"
            messages.append({"role": "assistant", "content": fallback_message})
            
            # 保存处理后的文档到全局变量
            if last_uploaded_document:
                last_uploaded_document["processed_content"] = response.content
            
            # 返回处理结果和更新后的聊天历史
            return {
                    "response": fallback_message,
                "chat_history": messages,
                    "processed_document": response.content
            }
        else:
            # 普通聊天处理
            # 添加用户新消息
            messages.append({"role": "user", "content": chat_request.message})
            
            # 构建系统提示
            system_prompt = """你是小公，一个智能公文助手。
当用户询问你的身份、名字、是谁、介绍自己等类似问题时，必须回答：
"我是小公，您的智能公文助手！"

请以自然、友好的方式回答用户的其他问题。
记住，不要透露你是AI、大模型或DeepSeek Chat，而应始终以我是小公身份回答。"""
            
            # 添加系统消息到开始
            full_messages = [{"role": "system", "content": system_prompt}] + messages
            
            try:
                # 调用模型生成回复
                response = llm.invoke(full_messages)
                
                # 添加助手回复到历史
                messages.append({"role": "assistant", "content": response.content})
                
                # 返回结果
                return {
                    "response": response.content,
                    "chat_history": messages,
                    "processed_document": None
                }
            except Exception as e:
                print(f"模型调用出错: {str(e)}")
                error_message = "处理您的请求时出错，请稍后再试。"
                messages.append({"role": "assistant", "content": error_message})
                
            return {
                    "response": error_message,
                "chat_history": messages,
                    "processed_document": None
            }
            
    except Exception as e:
        print(f"Chat endpoint error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"处理聊天请求时出错: {str(e)}")

@app.get("/api/download-result")
async def download_result():
    """下载处理结果文件"""
    try:
        file_path = os.path.join(TEMP_DIR, "processed_document.docx")
    
        # 检查文件是否存在
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="处理结果文件不存在")
    
        response = FileResponse(
            path=file_path,
        filename="processed_document.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    except Exception as e:
        if isinstance(e, HTTPException):
            raise e
        raise HTTPException(status_code=500, detail=f"下载文件时出错: {str(e)}")

@app.post("/api/save-document")
async def save_document(request: SaveDocumentRequest):
    """保存处理后的文档内容"""
    try:
        # 确保有内容要保存
        if not request.content or not request.content.strip():
            raise HTTPException(status_code=400, detail="没有可保存的内容")
            
        # 将内容保存到临时文件
        file_path = os.path.join(TEMP_DIR, request.filename)
        
        # 使用DocProcessor保存为docx
        DocProcessor.save_text_to_docx(request.content, file_path)
        
        return {"status": "success", "message": "文档已保存", "path": file_path}
        
    except Exception as e:
        if isinstance(e, HTTPException):
            raise e
        raise HTTPException(status_code=500, detail=f"保存文档时出错: {str(e)}")

# 添加流式聊天接口
@app.get("/api/stream-chat")
async def stream_chat(request: Request, message: str, document_content: str = None, chat_history: str = "[]", file_id: str = None):
    """处理聊天请求，返回流式响应"""
    async def event_generator():
        try:
            print(f"收到流式聊天请求: {message[:50]}..." if len(message) > 50 else f"收到流式聊天请求: {message}")
            
            # 解析聊天历史
            history = json.loads(chat_history) if chat_history else []
            print(f"解析到聊天历史消息数: {len(history)}")
            
            # 准备聊天消息
            messages = history.copy()
            
            # 检查是否有文档内容，如果没有，尝试从last_uploaded_document获取
            global last_uploaded_document, forced_correction_mode, file_list
            has_document = False
            input_text_for_correction = None
            
            # 优先使用传递的document_content参数
            if document_content and document_content.strip():
                processed_doc_content = document_content
                has_document = True
                print(f"使用传递的document_content，长度: {len(processed_doc_content)}")
            # 如果有file_id参数，通过ID查找文件
            elif file_id:
                print(f"通过file_id查找文档: {file_id}")
                # 在file_list中查找匹配的文件
                matching_file = next((f for f in file_list if str(f["id"]) == file_id), None)
                
                if matching_file and last_uploaded_document and last_uploaded_document.get("document_id") == file_id:
                    processed_doc_content = last_uploaded_document.get("content", "")
                    has_document = processed_doc_content and processed_doc_content.strip() != ""
                    print(f"通过file_id找到文档，长度: {len(processed_doc_content) if has_document else 0}")
                else:
                    print(f"找不到匹配file_id的文档内容: {file_id}")
                    processed_doc_content = ""
            # 否则检查全局变量
            elif last_uploaded_document:
                processed_doc_content = last_uploaded_document.get("content", "")
                has_document = processed_doc_content and processed_doc_content.strip() != ""
                print(f"使用last_uploaded_document中的内容，长度: {len(processed_doc_content) if has_document else 0}")
            else:
                processed_doc_content = ""
                print("没有找到文档内容")
            
            # 获取客户端IP地址作为会话ID
            client_ip = request.client.host
            session_id = f"session_{client_ip}"
            
            # 进行意图分类
            intent = IntentClassifier.classify(message)
            print(f"聊天意图分类结果: {intent}")
            
            # 检查是否处于强制公文纠错模式
            is_forced_mode = forced_correction_mode.get(session_id, False)
            
            # 初始化会话计数器
            if session_id not in chat_counter:
                chat_counter[session_id] = 0
            
            # 如果是聊天意图，增加计数
            if intent == "chat":
                chat_counter[session_id] += 1
                print(f"闲聊计数增加: {session_id} → {chat_counter[session_id]}")
                
                # 检查是否达到闲聊次数限制
                if chat_counter[session_id] > 3:
                    # 进入强制公文纠错模式
                    forced_correction_mode[session_id] = True
                    is_forced_mode = True
                    print(f"进入强制公文纠错模式: {session_id}")
            
            # 如果是公文纠错相关意图，重置闲聊计数
            if intent in ["correction", "writing", "recorrection"]:
                chat_counter[session_id] = 0
                # 如果之前处于强制模式，现在解除
                if is_forced_mode:
                    forced_correction_mode[session_id] = False
                    print(f"退出强制公文纠错模式: {session_id}")
            
            # 如果处于强制公文纠错模式，且不是公文纠错相关意图，强制拒绝回答
            if is_forced_mode and intent == "chat":
                guidance_message = "您已经闲聊超过3次，我只能回答公文纠错相关的问题了。请输入与公文处理相关的问题，例如「如何纠错」、「如何润色文档」或直接要求我「对文档进行公文纠错」。"
                yield f"data: {json.dumps({'content': guidance_message})}\n\n"
                yield "data: end\n\n"
                return
            
            # 根据意图处理请求
            if intent in ["correction", "writing", "recorrection", "forced_correction"]:
                # 检查消息中是否包含需要直接纠错的文本
                # 检查常见的模式，如"对下面的句子进行公文纠错："、"帮我纠错："等
                correction_patterns = [
                    r"对.*句子.*纠错[:：](.+)",
                    r"纠错[:：](.+)",
                    r"润色[:：](.+)",
                    r"修改[:：](.+)",
                    r"公文纠错[:：](.+)",
                    r"请.*公文纠错[:：](.+)",  # 添加新模式匹配"请进行公文纠错："
                    r"进行.*纠错[:：](.+)",     # 添加新模式匹配"进行公文纠错："
                    r"请.*纠错[:：](.+)"        # 添加新模式匹配"请纠错："
                ]
                
                # 添加详细日志，记录每个模式的匹配结果
                print(f"尝试匹配文本提取模式，原始消息: '{message}'")
                extracted_text = None
                
                for pattern in correction_patterns:
                    match = re.search(pattern, message, re.DOTALL)
                    if match:
                        extracted_text = match.group(1).strip()
                        print(f"✓ 成功匹配模式 '{pattern}'，提取文本: '{extracted_text}'")
                        input_text_for_correction = extracted_text
                        break
                    else:
                        print(f"✗ 模式 '{pattern}' 不匹配")
                
                # 如果没有从模式中提取到文本，检查是否应该使用整个消息
                if not input_text_for_correction and not has_document and len(message) > 15 and "上传" not in message and "文档" not in message:
                    # 直接使用整个消息作为要纠错的内容
                    input_text_for_correction = message
                    print(f"没有匹配到特定模式，使用整个消息作为纠错内容: '{input_text_for_correction}'")
                
                # 情况1: 用户提供了要直接纠错的文本
                if input_text_for_correction:
                    # 发送处理状态
                    yield f"data: {json.dumps({'content': '正在处理您提供的文本...'})}\n\n"
                    
                    # 使用MCP CSC服务处理文本
                    print("调用MCP CSC服务处理文本...")
                    try:
                        csc_result = mcp_csc(input_text_for_correction)
                        print(f"MCP CSC服务返回结果: {csc_result.get('status', 'unknown')}")
                        
                        if csc_result["status"] == "success":
                            if csc_result.get("modified", False):
                                corrected_text = csc_result["corrected_text"]
                                result_message = f"已完成文本纠错。修改如下：\n\n{corrected_text}\n\n修改说明：{csc_result.get('changes', '未提供详细修改说明')}"
                            else:
                                corrected_text = input_text_for_correction
                                result_message = csc_result.get("message", "文本未发现需要纠错的内容，已保持原文。")
                            print("MCP服务处理文本成功")
                        else:
                            raise Exception(f"MCP服务返回错误: {csc_result.get('message', '未知错误')}")
                    except Exception as e:
                        # MCP服务失败，回退到LLM处理
                        print(f"MCP服务处理失败: {str(e)}，回退到DeepSeek LLM处理")
                        
                        # 使用LLM处理方式
                        system_prompt = "你是一个专业的公文写作助手，请对下面的文本进行公文纠错和润色，改善语言表达，使其更符合公文规范。"
                        
                        doc_messages = [
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": f"请对以下文本进行公文纠错:\n\n{input_text_for_correction}"}
                        ]
                        
                        try:
                            print("调用DeepSeek LLM处理文本...")
                            response = llm.invoke(doc_messages)
                            corrected_text = response.content
                            print("LLM处理文本成功")
                            result_message = f"已使用DeepSeek AI大模型处理您的文本。\n\n{corrected_text}"
                        except Exception as llm_error:
                            print(f"LLM调用失败: {str(llm_error)}")
                            corrected_text = input_text_for_correction
                            result_message = f"处理失败: {str(llm_error)}。返回原始文本。"
                    
                    yield f"data: {json.dumps({'content': result_message})}\n\n"
                    
                # 情况2: 需要处理上传的文档
                elif not has_document:
                    # 文档内容缺失，提示用户上传文档
                    content_message = '请先上传Word文档再进行完整文档的公文纠错。如果您只想纠正特定句子或段落，请直接输入要纠正的内容，例如"纠错：这是一段需要纠错的文字"。'
                    yield f"data: {json.dumps({'content': content_message})}\n\n"
                    yield "data: end\n\n"
                    return
                else:
                    # 发送处理状态
                    yield f"data: {json.dumps({'content': '正在处理文档...'})}\n\n"
                    
                    # 使用MCP CSC服务处理文档
                    print("调用MCP CSC服务处理文档...")
                    try:
                        csc_result = mcp_csc(processed_doc_content)
                        print(f"MCP CSC服务返回结果: {csc_result.get('status', 'unknown')}")
                        
                        if csc_result["status"] == "success":
                            if csc_result.get("modified", False):
                                processed_doc = csc_result["corrected_text"]
                                processing_message = f"已根据MCP协议完成文档纠错，共进行了以下修改：\n{csc_result.get('changes', '未提供详细修改说明')}"
                            else:
                                processed_doc = processed_doc_content
                                processing_message = "已完成文档检查，" + csc_result.get("message", "未发现需要纠错的内容，已保持原文。")
                            
                            print("MCP服务处理文档成功")
                        else:
                            raise Exception(f"MCP服务返回错误: {csc_result.get('message', '未知错误')}")
                    except Exception as e:
                        # MCP服务失败，回退到LLM处理
                        print(f"MCP服务处理失败: {str(e)}，回退到DeepSeek LLM处理")
                        
                        # 使用LLM处理方式
                        system_prompt = "你是一个专业的文档处理助手。请根据用户的指令处理以下文档内容。"
                        full_prompt = f"文档内容:\n{processed_doc_content}\n\n用户指令:\n{message}"
                        
                        doc_messages = [
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": full_prompt}
                        ]
                        
                        try:
                            print("调用DeepSeek LLM处理文档...")
                            response = llm.invoke(doc_messages)
                            processed_doc = response.content
                            print("LLM处理文档成功")
                            
                            # 保存处理后的文档内容
                            if last_uploaded_document:
                                last_uploaded_document["processed_content"] = processed_doc
                            
                            # 发送处理结果
                            processing_message = "已使用DeepSeek AI大模型处理您的文档。处理结果已显示在中间区域。"
                        except Exception as llm_error:
                            print(f"LLM调用失败: {str(llm_error)}")
                            processed_doc = processed_doc_content
                            processing_message = f"处理失败: {str(llm_error)}。返回原始文档内容。"
                    
                    # 保存处理后的文档内容
                    if last_uploaded_document:
                        last_uploaded_document["processed_content"] = processed_doc
                    
                    # 发送处理结果
                    yield f"data: {json.dumps({'content': processing_message, 'processed_document': processed_doc})}\n\n"
            else:
                # 普通聊天处理
                print("处理普通聊天请求...")
                # 添加用户新消息
                messages.append({"role": "user", "content": message})
                
                # 构建系统提示
                system_prompt = """你是小公，一个智能公文助手。
当用户询问你的身份、名字、是谁、介绍自己等类似问题时，必须回答：
"我是小公，您的智能公文助手！"

请以自然、友好的方式回答用户的其他问题。
记住，不要透露你是AI、大模型或DeepSeek Chat，而应始终以我是小公身份回答。"""
                
                # 添加系统消息到开始
                full_messages = [{"role": "system", "content": system_prompt}] + messages
                
                # 调用模型
                try:
                    print("调用DeepSeek LLM生成聊天回复...")
                    response = llm.invoke(full_messages)
                    assistant_response = response.content
                    print(f"LLM生成回复成功: {assistant_response[:50]}..." if len(assistant_response) > 50 else f"LLM生成回复成功: {assistant_response}")
                    
                    # 模拟流式响应，将回复分成多个部分发送
                    words = assistant_response.split()
                    chunk_size = max(1, len(words) // 10)  # 分成大约10个部分发送
                    
                    for i in range(0, len(words), chunk_size):
                        chunk = " ".join(words[i:i+chunk_size])
                        yield f"data: {json.dumps({'content': chunk})}\n\n"
                        await asyncio.sleep(0.1)  # 轻微延迟模拟打字效果
                except Exception as e:
                    print(f"调用LLM出错: {str(e)}")
                    error_message = f"处理您的请求时出错: {str(e)}。请稍后再试。"
                    yield f"data: {json.dumps({'content': error_message})}\n\n"
            
            # 发送结束信号
            yield "data: end\n\n"
            
        except Exception as e:
            print(f"流式聊天处理出错: {str(e)}")
            import traceback
            print(traceback.format_exc())
            yield f"data: {json.dumps({'content': f'处理请求时出错: {str(e)}'})}\n\n"
            yield "data: end\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream")

# 添加前端路由适配
@app.post("/file/upload")
async def file_upload_adapter(file: UploadFile = File(...)):
    """适配前端/file/upload请求，存储文件并返回标准格式响应"""
    try:
        global last_uploaded_document, file_list
        
        if not file.filename.endswith(('.docx')):
            raise HTTPException(status_code=400, detail="只支持.docx格式文件")
        
        # 读取文件内容到内存
        file_content = await file.read()
        
        # 使用DocProcessor处理文档
        content = DocProcessor.load_doc_stream(file_content, file.filename)
        
        # 生成唯一文档ID
        document_id = str(hash(content + file.filename))
        
        # 重要：保存最近上传的文档（用于公文纠错）
        # 这确保无论通过哪个接口上传，document_content都能被正确设置
        last_uploaded_document = {
            "document_id": document_id,
            "content": content,
            "original_content": content,
            "filename": file.filename,
            "message": "文档上传成功，请在聊天框中输入处理指令"
        }
        print(f"文档已上传并保存到last_uploaded_document，长度: {len(content)}")
        
        # 创建文件记录并添加到文件列表
        file_item = {
            "id": document_id,
            "filename": file.filename,
            "fileuuid": document_id,
            "fileurl": "",
            "updatetime": datetime.now().isoformat(),
            "userid": "system"  # 可以从请求中获取，这里简化处理
        }
        
        # 将文件添加到文件列表
        if file_item not in file_list:
            file_list.append(file_item)
        
        # 返回符合前端期望的响应格式
        return {
            "code": 200,
            "data": file_item,
            "msg": "文件上传成功"
        }
    except Exception as e:
        print(f"文件上传处理出错: {str(e)}")
        return {
            "code": 500,
            "data": None,
            "msg": f"文件上传处理出错: {str(e)}"
        }

@app.post("/file/list")
async def file_list_adapter(request: Request):
    """适配前端/file/list请求，返回文件列表"""
    try:
        body = await request.json()
        # 返回文件列表
        return {
            "code": 200,
            "data": file_list,
            "msg": "获取文件列表成功"
        }
    except Exception as e:
        print(f"获取文件列表出错: {str(e)}")
        return {
            "code": 500,
            "data": [],
            "msg": f"获取文件列表出错: {str(e)}"
        }

@app.post("/chat/list")
async def chat_list_adapter(request: Request):
    """适配前端/chat/list请求"""
    try:
        # 尝试解析请求体，但不强制要求
        try:
            body = await request.json()
            print(f"收到chat/list请求，参数: {body}")
            # 根据userid过滤聊天列表
            filtered_chats = chat_list
            if body and 'userid' in body:
                userid = body.get('userid')
                filtered_chats = [chat for chat in chat_list if chat.get('userid') == userid]
        except json.JSONDecodeError:
            # 如果请求体为空或无效JSON，使用空字典
            body = {}
            print(f"收到chat/list请求，无参数")
            filtered_chats = chat_list
        
        # 返回过滤后的聊天列表
        response = {
            "code": 200,
            "data": filtered_chats,
            "msg": "获取聊天列表成功"
        }
        print(f"返回聊天列表响应: {response}")
        return response
    except Exception as e:
        print(f"获取聊天列表出错: {str(e)}")
        return {
            "code": 500,
            "data": [],
            "msg": f"获取聊天列表出错: {str(e)}"
        }

@app.post("/chat/create")
async def chat_create_adapter(request: Request):
    """适配前端/chat/create请求"""
    try:
        # 解析请求体
        try:
            body = await request.json()
            print(f"收到chat/create请求，参数: {body}")
        except json.JSONDecodeError:
            print("收到chat/create请求，但请求体不是有效的JSON")
            return {
                "code": 400,
                "data": None,
                "msg": "请求体格式不正确"
            }
        
        # 将聊天数据添加到列表中
        chat_id = body.get('id')
        # 检查是否已存在相同ID的聊天
        existing_chat_index = next((i for i, chat in enumerate(chat_list) if chat.get('id') == chat_id), -1)
        if existing_chat_index != -1:
            # 更新现有聊天
            chat_list[existing_chat_index] = body
        else:
            # 添加新聊天
            chat_list.append(body)
        
        # 保存聊天历史到文件
        save_chat_history()
        
        print(f"当前聊天列表长度: {len(chat_list)}")
        
        # 直接返回传入的数据，表示创建成功
        response = {
            "code": 200,
            "data": body,
            "msg": "创建聊天成功"
        }
        print(f"返回chat/create响应: {response}")
        return response
    except Exception as e:
        error_msg = f"创建聊天记录出错: {str(e)}"
        print(error_msg)
        import traceback
        print(traceback.format_exc())
        return {
            "code": 500,
            "data": None,
            "msg": error_msg
        }

@app.post("/chat/update")
async def chat_update_adapter(request: Request):
    """适配前端/chat/update请求"""
    try:
        # 解析请求体
        try:
            body = await request.json()
            print(f"收到chat/update请求，参数: {body}")
        except json.JSONDecodeError:
            print("收到chat/update请求，但请求体不是有效的JSON")
            return {
                "code": 400,
                "data": None,
                "msg": "请求体格式不正确"
            }
        
        # 更新聊天列表中的数据
        chat_id = body.get('id')
        if chat_id:
            existing_chat_index = next((i for i, chat in enumerate(chat_list) if chat.get('id') == chat_id), -1)
            if existing_chat_index != -1:
                # 更新现有聊天
                chat_list[existing_chat_index] = body
                print(f"已更新ID为{chat_id}的聊天")
            else:
                # 如果找不到，添加为新聊天
                chat_list.append(body)
                print(f"未找到ID为{chat_id}的聊天，已添加为新聊天")
            
            # 保存聊天历史到文件
            save_chat_history()
        
        # 直接返回传入的数据，表示更新成功
        return {
            "code": 200,
            "data": body,
            "msg": "更新聊天成功"
        }
    except Exception as e:
        print(f"更新聊天记录出错: {str(e)}")
        return {
            "code": 500,
            "data": None,
            "msg": f"更新聊天记录出错: {str(e)}"
        }

@app.delete("/chat/delete/{id}")
async def chat_delete_adapter(id: str):
    """适配前端/chat/delete/{id}请求"""
    try:
        # 查找并删除指定ID的聊天
        global chat_list
        original_length = len(chat_list)
        chat_list = [chat for chat in chat_list if chat.get('id') != id]
        
        if len(chat_list) < original_length:
            print(f"已删除ID为{id}的聊天")
            # 保存聊天历史到文件
            save_chat_history()
        else:
            print(f"未找到ID为{id}的聊天")
        
        # 返回成功响应
        return {
            "code": 200,
            "data": None,
            "msg": f"聊天{id}删除成功"
        }
    except Exception as e:
        print(f"删除聊天记录出错: {str(e)}")
        return {
            "code": 500,
            "data": None,
            "msg": f"删除聊天记录出错: {str(e)}"
        }

@app.post("/file/history")
async def file_history_adapter(request: Request):
    """获取文件历史记录"""
    try:
        try:
            body = await request.json()
            print(f"收到file/history请求，参数: {body}")
            # 根据userid过滤文件历史记录
            filtered_history = file_history_list
            if body and 'userid' in body:
                userid = body.get('userid')
                filtered_history = [file for file in file_history_list if file.get('userid') == userid]
        except json.JSONDecodeError:
            # 如果请求体为空或无效JSON，使用空字典
            body = {}
            print(f"收到file/history请求，无参数")
            filtered_history = file_history_list
        
        # 返回过滤后的文件历史记录
        response = {
            "code": 200,
            "data": filtered_history,
            "msg": "获取文件历史记录成功"
        }
        print(f"返回文件历史记录响应: {response}")
        return response
    except Exception as e:
        print(f"获取文件历史记录出错: {str(e)}")
        return {
            "code": 500,
            "data": [],
            "msg": f"获取文件历史记录出错: {str(e)}"
        }

@app.post("/file/add-history")
async def add_to_history_adapter(request: Request):
    """添加文件到历史记录"""
    try:
        # 解析请求体
        try:
            body = await request.json()
            print(f"收到file/add-history请求，参数: {body}")
        except json.JSONDecodeError:
            print("收到file/add-history请求，但请求体不是有效的JSON")
            return {
                "code": 400,
                "data": None,
                "msg": "请求体格式不正确"
            }
        
        # 检查必要字段
        file_id = body.get('id')
        if not file_id:
            return {
                "code": 400,
                "data": None,
                "msg": "缺少必要的文件ID字段"
            }
        
        # 查找是否已存在该文件的历史记录
        global file_history_list
        existing_index = next((i for i, file in enumerate(file_history_list) if file.get('id') == file_id), -1)
        
        # 设置访问次数和最近访问时间
        if 'accessCount' not in body:
            body['accessCount'] = 1
        if 'lastAccessed' not in body:
            body['lastAccessed'] = datetime.now().isoformat()
        
        if existing_index != -1:
            # 更新现有记录
            file_history_list[existing_index] = body
            print(f"更新ID为{file_id}的文件历史记录")
        else:
            # 添加新记录
            file_history_list.append(body)
            print(f"添加ID为{file_id}的文件历史记录")
        
        # 保存到文件
        save_file_history()
        
        # 返回更新后的文件记录
        return {
            "code": 200,
            "data": body,
            "msg": "添加文件历史记录成功"
        }
    except Exception as e:
        print(f"添加文件历史记录出错: {str(e)}")
        return {
            "code": 500,
            "data": None,
            "msg": f"添加文件历史记录出错: {str(e)}"
        }

@app.delete("/file/remove-history/{id}")
async def remove_from_history_adapter(id: str):
    """从历史记录中删除文件"""
    try:
        # 查找并删除指定ID的历史记录
        global file_history_list
        original_length = len(file_history_list)
        file_history_list = [file for file in file_history_list if file.get('id') != id]
        
        if len(file_history_list) < original_length:
            print(f"已从历史记录中删除ID为{id}的文件")
            # 保存到文件
            save_file_history()
        else:
            print(f"未找到ID为{id}的文件历史记录")
        
        # 返回成功响应
        return {
            "code": 200,
            "data": None,
            "msg": f"文件{id}历史记录删除成功"
        }
    except Exception as e:
        print(f"删除文件历史记录出错: {str(e)}")
        return {
            "code": 500,
            "data": None,
            "msg": f"删除文件历史记录出错: {str(e)}"
        }

@app.post("/file/clear-history")
async def clear_history_adapter(request: Request):
    """清空指定用户的文件历史记录"""
    try:
        # 解析请求体
        try:
            body = await request.json()
            print(f"收到file/clear-history请求，参数: {body}")
        except json.JSONDecodeError:
            print("收到file/clear-history请求，但请求体不是有效的JSON")
            return {
                "code": 400,
                "data": None,
                "msg": "请求体格式不正确"
            }
        
        # 获取用户ID
        userid = body.get('userid')
        if not userid:
            return {
                "code": 400,
                "data": None,
                "msg": "缺少必要的用户ID字段"
            }
        
        # 删除该用户的所有历史记录
        global file_history_list
        original_length = len(file_history_list)
        file_history_list = [file for file in file_history_list if file.get('userid') != userid]
        
        if len(file_history_list) < original_length:
            print(f"已清空用户{userid}的历史记录")
            # 保存到文件
            save_file_history()
        else:
            print(f"未找到用户{userid}的历史记录")
        
        # 返回成功响应
        return {
            "code": 200,
            "data": None,
            "msg": f"用户{userid}的历史记录已清空"
        }
    except Exception as e:
        print(f"清空历史记录出错: {str(e)}")
        return {
            "code": 500,
            "data": None,
            "msg": f"清空历史记录出错: {str(e)}"
        }

@app.delete("/file/delete/{id}")
async def file_delete_adapter(id: str):
    """适配前端文件删除请求"""
    # 简单返回成功，可根据需要扩展实际删除逻辑
    return {
        "code": 200,
        "data": None,
        "msg": f"文件{id}删除成功"
    }

@app.get("/file/preview")
async def file_preview_adapter(id: str):
    """适配前端文件预览请求"""
    try:
        # 从最近上传的文档内容返回
        if not last_uploaded_document:
            raise HTTPException(status_code=404, detail="无法预览文件，找不到内容")
        
        # 获取处理后的内容，如果存在的话
        content = last_uploaded_document.get("processed_content", last_uploaded_document["content"])
        
        # 将文本内容转换为docx并返回
        doc_bytes = convert_text_to_docx_bytes(content)
        
        return Response(
            content=doc_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"inline; filename=preview_{last_uploaded_document['filename']}"}
        )
    except Exception as e:
        print(f"预览文件出错: {str(e)}")
        raise HTTPException(status_code=500, detail=f"预览文件出错: {str(e)}")

@app.get("/file/download")
async def file_download_adapter(id: str):
    """适配前端文件下载请求"""
    try:
        # 从最近上传的文档内容返回
        if not last_uploaded_document:
            raise HTTPException(status_code=404, detail="无法下载文件，找不到内容")
        
        # 获取处理后的内容，如果存在的话
        content = last_uploaded_document.get("processed_content", last_uploaded_document["content"])
        
        # 生成时间戳
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        filename = last_uploaded_document["filename"]
        filename_base = filename.rsplit(".", 1)[0]
        download_filename = f"{filename_base}_processed_{timestamp}.docx"
        
        # 将文本内容转换为docx并返回
        doc_bytes = convert_text_to_docx_bytes(content)
        
        return Response(
            content=doc_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={download_filename}"}
        )
    except Exception as e:
        print(f"下载文件出错: {str(e)}")
        raise HTTPException(status_code=500, detail=f"下载文件出错: {str(e)}")

@app.post("/file/modified-preview")
async def file_modified_preview_adapter(request: Request):
    """适配前端修改后文件预览请求"""
    try:
        body = await request.json()
        file = body.get("file", {})
        chat = body.get("chat", {})
        
        # 可以根据需要构建特定的预览文档
        # 简单起见，直接返回一个示例文档
        sample_content = "这是修改后的文档内容示例"
        doc_bytes = DocProcessor.convert_text_to_docx_bytes(sample_content)
        
        return Response(
            content=doc_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"inline; filename=modified_preview.docx"}
        )
    except Exception as e:
        print(f"预览修改后文件出错: {str(e)}")
        raise HTTPException(status_code=500, detail=f"预览修改后文件出错: {str(e)}")

@app.post("/file/confirm-modify")
async def file_confirm_modify_adapter(request: Request):
    """适配前端确认修改请求"""
    try:
        body = await request.json()
        # 简单返回成功，可以扩展实际处理逻辑
        return {
            "code": 200,
            "data": None,
            "msg": "确认修改成功"
        }
    except Exception as e:
        print(f"确认修改出错: {str(e)}")
        raise HTTPException(status_code=500, detail=f"确认修改出错: {str(e)}")

# 将文本转换为Word文档
def convert_text_to_docx_bytes(text_content):
    """将文本内容转换为Word文档字节流"""
    try:
        # 创建一个新的Word文档
        doc = Document()
        
        # 按段落拆分文本并添加到文档中
        paragraphs = text_content.split('\n')
        for para in paragraphs:
            if para.strip():  # 跳过空行
                doc.add_paragraph(para)
        
        # 将文档保存到内存中
        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        # 返回字节内容
        return docx_bytes.getvalue()
    except Exception as e:
        logging.error(f"Convert text to DOCX error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to convert text to DOCX: {str(e)}")

# 添加到DocProcessor类
if not hasattr(DocProcessor, 'convert_text_to_docx_bytes'):
    DocProcessor.convert_text_to_docx_bytes = staticmethod(convert_text_to_docx_bytes)

# 添加适配格式的聊天API
@app.post("/chat/send")
async def chat_send_adapter(request: Request):
    """适配前端/chat/send请求，处理聊天和公文纠错"""
    try:
        body = await request.json()
        message = body.get("message", "")
        print(f"收到聊天请求: {message[:50]}..." if len(message) > 50 else f"收到聊天请求: {message}")
        
        # 获取聊天历史
        chat_history = body.get("chat_history", [])
        
        # 从最近上传的文档获取内容
        document_content = None
        if last_uploaded_document:
            document_content = last_uploaded_document["content"]
            print(f"已获取文档内容，长度: {len(document_content)}")
        else:
            print("无上传文档")
        
        # 进行意图分类
        intent = IntentClassifier.classify(message)
        print(f"聊天意图分类结果: {intent}")
        
        # 获取客户端IP地址作为会话ID
        client_ip = request.client.host
        session_id = f"session_{client_ip}"
        
        # 初始化会话计数器
        if session_id not in chat_counter:
            chat_counter[session_id] = 0
        
        # 如果是聊天意图，增加计数
        if intent == "chat":
            chat_counter[session_id] += 1
            print(f"闲聊计数增加: {session_id} → {chat_counter[session_id]}")
        
        # 检查是否超过闲聊次数限制
        if intent == "chat" and chat_counter[session_id] > 3:
            # 重置意图为强制文档处理
            intent = "forced_correction"
            # 添加系统消息，引导用户回到公文纠错
            guidance_message = "您已经闲聊超过3次，请回到公文纠错相关任务。您可以尝试输入「纠错」、「润色」等指令来处理您的文档。"
            
            # 向聊天历史添加用户消息和助手回复
            chat_history.append({"role": "user", "content": message})
            chat_history.append({"role": "assistant", "content": guidance_message})
            
            print("已触发闲聊限制，引导用户回到公文纠错")
            
            return {
                "code": 200,
                "data": {
                    "response": guidance_message,
                    "chat_history": chat_history,
                    "processed_document": None
                },
                "msg": "聊天处理成功"
            }
        
        # 根据意图处理请求
        if intent in ["correction", "writing", "recorrection", "forced_correction"]:
            # 检查是否有文档内容
            if not document_content:
                # 文档内容缺失，提示用户上传文档
                no_doc_message = "请先上传Word文档再进行公文纠错。您可以点击左侧面板的「上传文档」按钮上传文件。"
                
                # 向聊天历史添加用户消息和助手回复
                chat_history.append({"role": "user", "content": message})
                chat_history.append({"role": "assistant", "content": no_doc_message})
                
                print("无文档内容，提示用户上传文档")
                
                return {
                    "code": 200,
                    "data": {
                        "response": no_doc_message,
                        "chat_history": chat_history,
                        "processed_document": None
                    },
                    "msg": "聊天处理成功"
                }
            
            print("开始处理文档纠错请求...")
            # 使用MCP CSC服务处理文档
            try:
                print("调用MCP CSC服务处理文档中...")
                csc_result = mcp_csc(document_content)
                print(f"MCP CSC服务返回结果: {csc_result.get('status', 'unknown')}")
                
                if csc_result["status"] == "success":
                    if csc_result.get("modified", False):
                        processed_document = csc_result["corrected_text"]
                        processing_message = f"已完成文档纠错，共进行了以下修改：\n{csc_result.get('changes', '未提供详细修改说明')}"
                    else:
                        processed_document = document_content
                        processing_message = "已完成文档检查，" + csc_result.get("message", "未发现需要纠错的内容，已保持原文。")
                    
                    print("MCP服务处理成功")
                else:
                    raise Exception(f"MCP服务返回错误: {csc_result.get('message', '未知错误')}")
            except Exception as e:
                print(f"MCP服务处理失败: {str(e)}，回退到LLM处理")
                
                # 使用LLM处理方式
                system_prompt = "你是一个专业的公文写作助手，请对下面的文本进行公文纠错和润色，改善语言表达，使其更符合公文规范。"
                
                doc_messages = [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"请对以下文本进行公文纠错:\n\n{document_content}"}
                ]
                
                try:
                    print("调用DeepSeek LLM处理文档...")
                    response = llm.invoke(doc_messages)
                    processed_document = response.content
                    print("LLM处理文档成功")
                    
                    # 保存处理后的文档内容
                    if last_uploaded_document:
                        last_uploaded_document["processed_content"] = processed_document
                    
                    # 发送处理结果
                    processing_message = "已使用DeepSeek AI大模型处理您的文档。处理结果已显示在中间区域。"
                except Exception as llm_error:
                    print(f"LLM调用失败: {str(llm_error)}")
                    processed_document = document_content
                    processing_message = f"处理失败: {str(llm_error)}。返回原始文档内容。"
            
            # 保存处理后的文档内容
            if last_uploaded_document:
                last_uploaded_document["processed_content"] = processed_document
            
            # 向聊天历史添加用户消息和助手回复
            chat_history.append({"role": "user", "content": message})
            chat_history.append({"role": "assistant", "content": processing_message})
            
            return {
                "code": 200,
                "data": {
                    "response": processing_message,
                    "chat_history": chat_history,
                    "processed_document": processed_document
                },
                "msg": "公文纠错处理成功"
            }
        else:
            print("处理普通聊天请求...")
            # 普通聊天处理
            # 构建系统提示
            system_prompt = """你是小公，一个智能公文助手。
当用户询问你的身份、名字、是谁、介绍自己等类似问题时，必须回答：
"我是小公，您的智能公文助手！"

请以自然、友好的方式回答用户的其他问题。
记住，不要透露你是AI、大模型或DeepSeek Chat，而应始终以我是小公身份回答。"""
            
            # 添加系统消息到开始
            full_messages = [{"role": "system", "content": system_prompt}]
            
            # 添加历史消息和当前用户消息
            for msg in chat_history:
                if msg["role"] in ["user", "assistant"]:
                    full_messages.append(msg)
            
            # 确保最新的用户消息被添加到messages中
            user_message_exists = False
            for msg in full_messages:
                if msg["role"] == "user" and msg["content"] == message:
                    user_message_exists = True
                    break
            
            if not user_message_exists:
                full_messages.append({"role": "user", "content": message})
            
            print(f"准备调用DeepSeek LLM进行聊天，消息数量: {len(full_messages)}")
            
            # 调用模型生成回复
            try:
                # 记录调用过程
                print("调用DeepSeek LLM生成回复...")
                model_response = llm.invoke(full_messages)
                assistant_message = model_response.content
                print(f"LLM生成回复成功: {assistant_message[:50]}..." if len(assistant_message) > 50 else f"LLM生成回复成功: {assistant_message}")
                
                # 添加用户消息和助手回复到历史
                chat_history.append({"role": "user", "content": message})
                chat_history.append({"role": "assistant", "content": assistant_message})
                
                return {
                    "code": 200,
                    "data": {
                        "response": assistant_message,
                        "chat_history": chat_history,
                        "processed_document": None
                    },
                    "msg": "聊天处理成功"
                }
            except Exception as e:
                print(f"调用LLM出错: {str(e)}")
                error_message = f"处理您的请求时出错: {str(e)}。请稍后再试。"
                chat_history.append({"role": "user", "content": message})
                chat_history.append({"role": "assistant", "content": error_message})
                
                return {
                    "code": 200,
                    "data": {
                        "response": error_message,
                        "chat_history": chat_history,
                        "processed_document": None
                    },
                    "msg": "聊天处理失败"
                }
    
    except Exception as e:
        print(f"处理聊天请求出错: {str(e)}")
        import traceback
        print(traceback.format_exc())
        return {
            "code": 500,
            "data": None,
            "msg": f"处理聊天请求出错: {str(e)}"
        }

# 添加MCP服务接口
@app.get("/api/mcp/info")
async def mcp_info():
    """返回MCP服务信息"""
    try:
        # 读取配置文件
        config_path = Path(__file__).parent / "mcp_config.json"
        with open(config_path, "r", encoding="utf-8") as f:
            config = json.load(f)
            
        return {
            "status": "success",
            "service_name": config["service"]["name"],
            "version": config["service"]["version"],
            "description": config["service"]["description"],
            "tools": [
                {"name": tool["name"], "description": tool["description"]}
                for tool in config["tools"]
            ]
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}

@app.post("/api/mcp/correction")
async def mcp_correction(request: Request):
    """MCP公文纠错接口"""
    try:
        body = await request.json()
        text = body.get("text")
        
        if not text:
            return {"status": "error", "message": "文本内容不能为空"}
            
        # 调用agent_tools中的CSC纠错方法
        result = mcp_csc(text)
        
        # 检查结果状态
        if result["status"] == "success":
            if result.get("modified", False):
                return {
                    "status": "success", 
                    "corrected": result["corrected_text"],
                    "changes": result.get("changes", "")
                }
            else:
                return {
                    "status": "success", 
                    "corrected": text,
                    "message": result.get("message", "未发现需要纠错的内容")
                }
        else:
            # 错误处理
            return {
                "status": "error", 
                "message": result.get("message", "纠错处理失败")
            }
            
    except Exception as e:
        import traceback
        error_detail = str(e) + "\n" + traceback.format_exc()
        logging.error(f"MCP纠错接口出错: {error_detail}")
        return {"status": "error", "message": str(e)}

@app.post("/api/mcp/invoke")
async def mcp_invoke_tool(request: Request):
    """通用MCP工具调用接口"""
    try:
        body = await request.json()
        tool = body.get("tool")
        params = body.get("params", {})
        
        if not tool:
            return {"status": "error", "message": "工具名称不能为空"}
            
        # 使用agent_tools中的call_mcp_service方法
        result = call_mcp_service(tool, params)
        
        return result
    except Exception as e:
        import traceback
        error_detail = str(e) + "\n" + traceback.format_exc()
        logging.error(f"MCP通用接口调用出错: {error_detail}")
        return {"status": "error", "message": str(e)}

@app.post("/api/convert-text-to-docx")
async def convert_text_to_docx(request: Request):
    """
    将文本内容转换为DOCX格式
    """
    try:
        data = await request.json()
        text_content = data.get("text_content", "")
        filename = data.get("filename", "converted_document")
        
        if not text_content:
            raise HTTPException(status_code=400, detail="No text content provided")
        
        # 创建一个新的Word文档
        doc = Document()
        
        # 按段落拆分文本并添加到文档中
        paragraphs = text_content.split('\n')
        for para in paragraphs:
            if para.strip():  # 跳过空行
                doc.add_paragraph(para)
        
        # 将文档保存到内存中
        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        # 如果文件名没有.docx后缀，添加它
        if not filename.lower().endswith('.docx'):
            filename = f"{filename}.docx"
        
        # 返回文档内容
        return StreamingResponse(
            docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except Exception as e:
        logging.error(f"Convert text to DOCX error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to convert text to DOCX: {str(e)}")

@app.post("/api/upload")
async def api_file_upload(file: UploadFile = File(...)):
    """处理前端/api/upload请求，转发到file/upload处理程序"""
    return await file_upload_adapter(file)

@app.post("/api/list")
async def api_file_list(request: Request):
    """处理前端/api/list请求，转发到file/list处理程序"""
    return await file_list_adapter(request)

@app.get("/api/preview")
async def api_file_preview(id: str):
    """处理前端/api/preview请求，转发到file/preview处理程序"""
    return await file_preview_adapter(id)

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("api.main:app", host="0.0.0.0", port=8003, reload=True) 